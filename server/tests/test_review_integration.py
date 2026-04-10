"""Integration tests for the review feature — mock Celery task via sys.modules."""
import sys
import uuid
import pytest
import pytest_asyncio
from unittest.mock import MagicMock
from docx import Document

from server.app.models.task import Task
from server.app.models.review_task import ReviewTask

pytestmark = pytest.mark.asyncio


def _install_review_task_mock():
    """Inject a fake review_task module so the lazy import in the router resolves."""
    fake_task = MagicMock()
    fake_task.delay.return_value = MagicMock(id="fake-review-celery-id")

    fake_module = MagicMock()
    fake_module.run_review = fake_task

    sys.modules["server.app.tasks.review_task"] = fake_module
    return fake_task


def _remove_review_task_mock():
    sys.modules.pop("server.app.tasks.review_task", None)


@pytest_asyncio.fixture
async def completed_task(db_session, test_user, tmp_path):
    """A completed bid task with extracted_data."""
    fpath = tmp_path / "招标文件.docx"
    doc = Document()
    doc.add_paragraph("测试招标文件内容")
    doc.save(str(fpath))
    task = Task(
        id=uuid.uuid4(), user_id=test_user.id,
        filename="招标文件.docx", file_path=str(fpath),
        file_size=1000, status="completed",
        extracted_data={
            "schema_version": "1.0",
            "modules": {
                "module_e": {
                    "sections": [{"id": "risks", "type": "table", "title": "废标风险",
                                  "columns": ["序号", "风险项", "依据"],
                                  "rows": [["1", "未密封", "投标须知"]]}]
                },
            },
        },
    )
    db_session.add(task)
    await db_session.commit()
    await db_session.refresh(task)
    return task


@pytest.fixture
def tender_docx(tmp_path):
    """A minimal tender docx for upload."""
    fpath = tmp_path / "投标文件.docx"
    doc = Document()
    doc.add_paragraph("第一章 投标函")
    doc.add_paragraph("我方承诺按照招标文件要求提交投标文件。")
    doc.save(str(fpath))
    return fpath


async def test_create_review(client, auth_headers, completed_task, tender_docx):
    """POST /api/reviews creates a review task."""
    fake_task = _install_review_task_mock()
    try:
        with open(tender_docx, "rb") as f:
            resp = await client.post(
                "/api/reviews",
                data={"bid_task_id": str(completed_task.id)},
                files={"tender_file": ("投标文件.docx", f, "application/octet-stream")},
                headers=auth_headers,
            )
        assert resp.status_code == 201
        data = resp.json()
        assert "id" in data
        assert data["version"] == 1
    finally:
        _remove_review_task_mock()


async def test_list_reviews_after_create(client, auth_headers, completed_task, tender_docx):
    """GET /api/reviews returns created reviews."""
    fake_task = _install_review_task_mock()
    try:
        with open(tender_docx, "rb") as f:
            await client.post(
                "/api/reviews",
                data={"bid_task_id": str(completed_task.id)},
                files={"tender_file": ("投标文件.docx", f, "application/octet-stream")},
                headers=auth_headers,
            )
        resp = await client.get("/api/reviews", headers=auth_headers)
        assert resp.status_code == 200
        assert resp.json()["total"] >= 1
    finally:
        _remove_review_task_mock()


async def test_delete_review_cleanup(client, auth_headers, completed_task, tender_docx):
    """DELETE /api/reviews/{id} removes files and DB record."""
    fake_task = _install_review_task_mock()
    try:
        with open(tender_docx, "rb") as f:
            create_resp = await client.post(
                "/api/reviews",
                data={"bid_task_id": str(completed_task.id)},
                files={"tender_file": ("投标文件.docx", f, "application/octet-stream")},
                headers=auth_headers,
            )
        review_id = create_resp.json()["id"]
        del_resp = await client.delete(f"/api/reviews/{review_id}", headers=auth_headers)
        assert del_resp.status_code == 204
        # Verify it's gone
        get_resp = await client.get(f"/api/reviews/{review_id}", headers=auth_headers)
        assert get_resp.status_code == 404
    finally:
        _remove_review_task_mock()
