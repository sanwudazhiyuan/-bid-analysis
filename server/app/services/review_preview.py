"""Build preview HTML for a completed review.

Extracted so the heavy python-docx rendering can run once at review
completion time (in the Celery worker) instead of on every preview request.
"""
from html import escape as html_escape


def build_preview_html(
    tender_file_path: str,
    review_items: list[dict],
    extracted_images: list[dict],
    review_id: str,
) -> str:
    from server.app.routers.files import _para_to_html, _table_to_html
    from docx import Document
    from docx.oxml.ns import qn

    para_review_map: dict[int, list[dict]] = {}
    for item in (review_items or []):
        if item.get("result") in ("fail", "warning"):
            for loc in item.get("tender_locations", []):
                for pi in loc.get("para_indices", []):
                    para_review_map.setdefault(pi, []).append(item)

    doc = Document(tender_file_path)
    body = doc.element.body
    parts: list[tuple[str, int]] = []
    element_idx = 0
    para_idx = 0
    table_idx = 0

    for child in body:
        if child.tag == qn("w:p"):
            if para_idx < len(doc.paragraphs):
                html_str = _para_to_html(doc.paragraphs[para_idx])
                if html_str and element_idx in para_review_map:
                    items = para_review_map[element_idx]
                    review_ids = " ".join(str(item["id"]) for item in items)
                    result = "fail" if any(i["result"] == "fail" for i in items) else "warning"
                    css_class = f"review-highlight review-{result}"
                    html_str = html_str.replace(
                        "<p", f'<p data-review-id="{review_ids}" class="{css_class}"', 1
                    )
                if html_str:
                    parts.append((html_str, element_idx))
                para_idx += 1
            element_idx += 1
        elif child.tag == qn("w:tbl"):
            if table_idx < len(doc.tables):
                parts.append((_table_to_html(doc.tables[table_idx]), element_idx))
                table_idx += 1
            element_idx += 1

    para_image_map: dict[int, list[str]] = {}
    for img in extracted_images or []:
        indices = img.get("near_para_indices")
        if not indices:
            pi = img.get("near_para_index")
            indices = [pi] if pi is not None else []
        for pi in indices:
            para_image_map.setdefault(pi, []).append(img["filename"])

    final_parts: list[str] = []
    for html_part, elem_idx in parts:
        final_parts.append(html_part)
        if elem_idx in para_image_map:
            for fn in para_image_map[elem_idx]:
                safe_fn = html_escape(fn)
                final_parts.append(
                    f'<div class="review-image" data-para-index="{elem_idx}">'
                    f'<img src="/api/reviews/{review_id}/images/{safe_fn}" '
                    f'alt="{safe_fn}" loading="lazy" />'
                    f'</div>'
                )

    return "\n".join(final_parts)
