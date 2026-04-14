# server/app/routers/files.py
"""Router for file management — list, download, preview, delete."""
import os
import html as _html
import uuid as _uuid

from fastapi import APIRouter, Depends, HTTPException, Query
from fastapi.responses import FileResponse
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from server.app.database import get_db
from server.app.deps import get_current_user
from server.app.models.user import User
from server.app.models.task import Task
from server.app.models.generated_file import GeneratedFile
from server.app.services.file_service import list_files, FILE_TYPE_MAP

router = APIRouter(prefix="/api/files", tags=["files"])


@router.get("")
async def list_files_endpoint(
    file_type: str = Query(..., description="bid-documents|reports|formats|checklists"),
    page: int = 1,
    page_size: int = 20,
    q: str | None = None,
    user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    if file_type not in FILE_TYPE_MAP:
        raise HTTPException(status_code=400, detail=f"Invalid file_type: {file_type}")
    items, total = await list_files(db, user.id, file_type, page, page_size, q)
    return {"items": items, "total": total, "page": page, "page_size": page_size}


@router.get("/{file_type}/{file_id}/download")
async def download_file(
    file_type: str, file_id: str,
    user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    if file_type == "bid-documents":
        task = await _get_user_task(db, file_id, user.id)
        if not task or not os.path.exists(task.file_path):
            raise HTTPException(status_code=404, detail="File not found")
        return FileResponse(
            task.file_path,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            filename=task.filename,
        )
    else:
        gf = await _get_generated_file_by_task(db, file_id, file_type, user.id)
        if not gf:
            gf = await _get_user_generated_file(db, file_id, user.id)
        if not gf or not os.path.exists(gf.file_path):
            raise HTTPException(status_code=404, detail="File not found")
        return FileResponse(
            gf.file_path,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            filename=os.path.basename(gf.file_path),
        )


@router.get("/{file_type}/{file_id}/preview")
async def preview_file(
    file_type: str, file_id: str,
    user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """Return HTML preview of a .docx file with formatting."""
    if file_type == "bid-documents":
        task = await _get_user_task(db, file_id, user.id)
        if not task or not os.path.exists(task.file_path):
            raise HTTPException(status_code=404, detail="File not found")
        file_path = task.file_path
        fname = task.filename
    else:
        gf = await _get_generated_file_by_task(db, file_id, file_type, user.id)
        if not gf:
            gf = await _get_user_generated_file(db, file_id, user.id)
        if not gf or not os.path.exists(gf.file_path):
            raise HTTPException(status_code=404, detail="File not found")
        file_path = gf.file_path
        fname = os.path.basename(gf.file_path)

    ext = os.path.splitext(file_path)[1].lower()
    if ext == ".docx":
        html_content = _docx_to_html(file_path)
    elif ext == ".doc":
        html_content = _doc_to_html(file_path)
    elif ext == ".pdf":
        html_content = _pdf_to_html(file_path)
    else:
        raise HTTPException(status_code=400, detail=f"不支持预览的文件类型: {ext}")

    return {"html": html_content, "filename": fname}


@router.delete("/{file_type}/{file_id}", status_code=204)
async def delete_file(
    file_type: str, file_id: str,
    user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    if file_type == "bid-documents":
        import shutil

        task = await _get_user_task(db, file_id, user.id)
        if not task:
            raise HTTPException(status_code=404, detail="Task not found")
        task_id_str = str(task.id)
        # Delete associated generated files from disk
        gen_result = await db.execute(
            select(GeneratedFile).where(GeneratedFile.task_id == task.id)
        )
        for gf in gen_result.scalars().all():
            if os.path.exists(gf.file_path):
                os.remove(gf.file_path)
            await db.delete(gf)
        # Delete associated review tasks (foreign key constraint)
        from server.app.models.review_task import ReviewTask
        review_result = await db.execute(
            select(ReviewTask).where(ReviewTask.bid_task_id == task.id)
        )
        for rt in review_result.scalars().all():
            # Delete review output files
            if rt.tender_file_path and os.path.exists(rt.tender_file_path):
                os.remove(rt.tender_file_path)
            if rt.annotated_file_path and os.path.exists(rt.annotated_file_path):
                os.remove(rt.annotated_file_path)
            await db.delete(rt)
        # Delete uploaded file
        if task.file_path and os.path.exists(task.file_path):
            os.remove(task.file_path)
        # Delete intermediate and output directories
        from server.app.config import settings as app_settings
        for subdir in ("intermediate", "output"):
            dirpath = os.path.join(app_settings.DATA_DIR, subdir, task_id_str)
            if os.path.isdir(dirpath):
                shutil.rmtree(dirpath, ignore_errors=True)
        await db.delete(task)
        await db.commit()
        return
    else:
        gf = await _get_generated_file_by_task(db, file_id, file_type, user.id)
        if not gf:
            gf = await _get_user_generated_file(db, file_id, user.id)
        if not gf:
            raise HTTPException(status_code=404, detail="File not found")
        if os.path.exists(gf.file_path):
            os.remove(gf.file_path)
        await db.delete(gf)
        await db.commit()


# ========== docx → HTML ==========


def _run_style(run) -> str:
    """Convert a single Run to HTML with inline styles."""
    text = _html.escape(run.text)
    if not text:
        return ""

    styles = []
    tags_open = []
    tags_close = []

    if run.bold:
        tags_open.append("<b>")
        tags_close.insert(0, "</b>")
    if run.italic:
        tags_open.append("<i>")
        tags_close.insert(0, "</i>")
    if run.underline:
        styles.append("text-decoration:underline")

    font = run.font
    if font.color and font.color.rgb:
        styles.append(f"color:#{font.color.rgb}")
    if font.size:
        # Convert EMU to pt
        pt = font.size.pt
        styles.append(f"font-size:{pt}pt")
    try:
        if font.highlight_color:
            styles.append("background-color:yellow")
    except ValueError:
        pass

    style_attr = f' style="{";".join(styles)}"' if styles else ""
    if style_attr and not tags_open:
        return f'<span{style_attr}>{text}</span>'
    elif tags_open:
        inner = f'<span{style_attr}>{text}</span>' if style_attr else text
        return "".join(tags_open) + inner + "".join(tags_close)
    return text


def _para_to_html(para) -> str:
    """Convert a Paragraph to HTML with alignment and heading styles."""
    style_name = para.style.name if para.style else ""
    runs_html = "".join(_run_style(r) for r in para.runs)

    # If no runs extracted, fall back to plain text
    if not runs_html and para.text:
        runs_html = _html.escape(para.text)

    if not runs_html:
        return ""

    # Alignment
    align = ""
    if para.alignment is not None:
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        align_map = {
            WD_ALIGN_PARAGRAPH.CENTER: "center",
            WD_ALIGN_PARAGRAPH.RIGHT: "right",
            WD_ALIGN_PARAGRAPH.JUSTIFY: "justify",
        }
        align = align_map.get(para.alignment, "")

    # Indentation
    indent_style = ""
    pf = para.paragraph_format
    if pf.first_line_indent and pf.first_line_indent.pt > 0:
        indent_style = f"text-indent:{pf.first_line_indent.pt}pt;"
    elif pf.left_indent and pf.left_indent.pt > 0:
        indent_style = f"margin-left:{pf.left_indent.pt}pt;"

    style_parts = []
    if align:
        style_parts.append(f"text-align:{align}")
    if indent_style:
        style_parts.append(indent_style)
    style_attr = f' style="{";".join(style_parts)}"' if style_parts else ""

    # Heading detection
    if "Heading" in style_name:
        level = style_name.replace("Heading", "").replace(" ", "").strip() or "3"
        try:
            level = str(min(int(level), 6))
        except ValueError:
            level = "3"
        return f"<h{level}{style_attr}>{runs_html}</h{level}>"

    # List detection
    if style_name.startswith("List"):
        return f"<li{style_attr}>{runs_html}</li>"

    return f"<p{style_attr}>{runs_html}</p>"


def _table_to_html(table) -> str:
    """Convert a Table to HTML with cell formatting."""
    rows_html = []
    for i, row in enumerate(table.rows):
        cells_html = []
        for cell in row.cells:
            cell_text = _html.escape(cell.text)
            tag = "th" if i == 0 else "td"
            style = "padding:6px 10px;border:1px solid #d1d5db;"
            if i == 0:
                style += "background:#f3f4f6;font-weight:600;"
            cells_html.append(f"<{tag} style=\"{style}\">{cell_text}</{tag}>")
        rows_html.append("<tr>" + "".join(cells_html) + "</tr>")
    return (
        '<table style="border-collapse:collapse;width:100%;margin:12px 0;font-size:14px">'
        + "".join(rows_html)
        + "</table>"
    )


def _docx_to_html(file_path: str) -> str:
    """Convert a .docx file to styled HTML."""
    from docx import Document

    doc = Document(file_path)
    parts = []

    # Interleave paragraphs and tables in document order via the XML body
    from docx.oxml.ns import qn
    body = doc.element.body
    para_idx = 0
    table_idx = 0

    for child in body:
        if child.tag == qn("w:p"):
            if para_idx < len(doc.paragraphs):
                html = _para_to_html(doc.paragraphs[para_idx])
                if html:
                    parts.append(html)
                para_idx += 1
        elif child.tag == qn("w:tbl"):
            if table_idx < len(doc.tables):
                parts.append(_table_to_html(doc.tables[table_idx]))
                table_idx += 1

    return "\n".join(parts)


def _doc_to_html(file_path: str) -> str:
    """Convert a .doc file to HTML via LibreOffice temp conversion."""
    import tempfile
    from src.parser.doc_parser import _convert_doc_to_docx

    with tempfile.TemporaryDirectory() as tmp_dir:
        docx_path = _convert_doc_to_docx(file_path, tmp_dir)
        return _docx_to_html(docx_path)


def _pdf_to_html(file_path: str) -> str:
    """Convert a PDF file to simple HTML with text and tables."""
    try:
        import pdfplumber
    except ImportError:
        return "<p>pdfplumber 未安装，无法预览 PDF</p>"

    parts = []
    with pdfplumber.open(file_path) as pdf:
        for page_num, page in enumerate(pdf.pages, 1):
            parts.append(
                f'<h2 style="color:#6b7280;font-size:12px;margin-top:20px;border-bottom:1px solid #e5e7eb;padding-bottom:4px">'
                f"第 {page_num} 页</h2>"
            )

            # Tables
            tables = page.extract_tables() or []
            for table in tables:
                if not table:
                    continue
                rows_html = []
                for i, row in enumerate(table):
                    cells = [
                        f'<td style="padding:4px 8px;border:1px solid #d1d5db">{_html.escape(cell or "")}</td>'
                        for cell in row
                    ]
                    rows_html.append(f"<tr>{''.join(cells)}</tr>")
                parts.append(
                    '<table style="border-collapse:collapse;width:100%;margin:8px 0;font-size:14px">'
                    + "".join(rows_html)
                    + "</table>"
                )

            # Text
            text = page.extract_text() or ""
            if text.strip():
                for line in text.split("\n"):
                    line = line.strip()
                    if line:
                        parts.append(
                            f'<p style="margin:2px 0;font-size:14px">{_html.escape(line)}</p>'
                        )

    return "\n".join(parts)


# ========== helpers ==========


async def _get_user_task(db: AsyncSession, task_id: str, user_id: int):
    try:
        task_uuid = _uuid.UUID(task_id)
    except ValueError:
        return None
    result = await db.execute(
        select(Task).where(Task.id == task_uuid, Task.user_id == user_id)
    )
    return result.scalar_one_or_none()


async def _get_user_generated_file(db: AsyncSession, file_id: str, user_id: int):
    """Look up generated file by GeneratedFile.id (int)."""
    try:
        fid = int(file_id)
        result = await db.execute(
            select(GeneratedFile)
            .join(Task, GeneratedFile.task_id == Task.id)
            .where(GeneratedFile.id == fid, Task.user_id == user_id)
        )
        found = result.scalar_one_or_none()
        if found:
            return found
    except (ValueError, TypeError):
        pass
    return None


async def _get_generated_file_by_task(db: AsyncSession, task_id: str, file_type: str, user_id: int):
    """Look up generated file by task UUID + file_type (report/format/checklist)."""
    db_type = {"reports": "report", "formats": "format", "checklists": "checklist"}.get(file_type)
    if not db_type:
        return None
    try:
        task_uuid = _uuid.UUID(task_id)
    except ValueError:
        return None
    result = await db.execute(
        select(GeneratedFile)
        .join(Task, GeneratedFile.task_id == Task.id)
        .where(GeneratedFile.task_id == task_uuid, GeneratedFile.file_type == db_type, Task.user_id == user_id)
    )
    return result.scalar_one_or_none()
