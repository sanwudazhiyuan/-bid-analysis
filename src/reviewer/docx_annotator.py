"""Generate annotated docx with summary table + highlights + Word comments."""
import os
import datetime
from lxml import etree
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Word XML namespaces
W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
R_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
COMMENTS_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
COMMENTS_CT = "application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml"
COMMENTS_REL = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments"

NSMAP = {"w": W_NS, "r": R_NS}


def _result_symbol(result: str) -> str:
    return {"pass": "✓", "fail": "✗", "warning": "⚠", "error": "?"}.get(result, "?")


def _result_color(result: str) -> RGBColor:
    return {
        "pass": RGBColor(0x22, 0x8B, 0x22),
        "fail": RGBColor(0xCC, 0x00, 0x00),
        "warning": RGBColor(0xFF, 0x8C, 0x00),
    }.get(result, RGBColor(0x80, 0x80, 0x80))


def _add_summary_section(doc: Document, review_items: list[dict], summary: dict,
                         bid_filename: str, tender_filename: str):
    """Insert summary header and table at the beginning of the document."""
    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("投标文件审查报告")
    run.bold = True
    run.font.size = Pt(18)

    # Meta info
    meta = doc.add_paragraph()
    meta.add_run(f"招标项目: {bid_filename}").font.size = Pt(10)
    meta.add_run("\n")
    meta.add_run(f"投标文件: {tender_filename}").font.size = Pt(10)
    meta.add_run("\n")
    meta.add_run(f"审查时间: {datetime.date.today().isoformat()}").font.size = Pt(10)

    # Stats
    stats = doc.add_paragraph()
    stats_text = f"共{summary['total']}条 | 通过{summary['pass']} | 不合规{summary['fail']} | 警告{summary['warning']}"
    if summary.get("critical_fails", 0) > 0:
        stats_text += f" | 废标风险: {summary['critical_fails']}条"
    stats.add_run(stats_text).font.size = Pt(10)

    # Summary table
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    headers = ["序号", "条款", "结果", "置信度", "说明"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(9)

    for idx, item in enumerate(review_items):
        row = table.add_row()
        row.cells[0].text = str(idx + 1)
        row.cells[1].text = item.get("clause_text", "")
        row.cells[2].text = _result_symbol(item.get("result", ""))
        row.cells[3].text = f"{item.get('confidence', 0)}%"
        row.cells[4].text = item.get("reason", "")
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(8)

    # Separator
    doc.add_paragraph("─" * 60)
    doc.add_paragraph()


class _CommentManager:
    """Manages Word comments XML in a docx document."""

    def __init__(self, doc: Document):
        self.doc = doc
        self._comments_element = None
        self._next_id = 0
        self._setup_comments_part()

    def _setup_comments_part(self):
        """Create or get the comments XML part in the docx package."""
        # Create comments XML root
        self._comments_element = etree.Element(f"{{{W_NS}}}comments", nsmap=NSMAP)

        # Add as a part to the document
        from docx.opc.part import Part
        from docx.opc.packuri import PackURI

        comments_xml = etree.tostring(self._comments_element, xml_declaration=True, encoding="UTF-8")
        part_name = PackURI("/word/comments.xml")

        doc_part = self.doc.part
        comments_part = Part(part_name, COMMENTS_CT, comments_xml, doc_part.package)
        doc_part.relate_to(comments_part, COMMENTS_REL)
        self._comments_part = comments_part

    def add_comment(self, para_element, comment_text: str, author: str = "AI审查") -> int:
        """Add a Word comment to a paragraph. Returns comment ID."""
        comment_id = str(self._next_id)
        self._next_id += 1

        # Add comment to comments.xml
        comment_elem = etree.SubElement(self._comments_element, f"{{{W_NS}}}comment")
        comment_elem.set(f"{{{W_NS}}}id", comment_id)
        comment_elem.set(f"{{{W_NS}}}author", author)
        comment_elem.set(f"{{{W_NS}}}date", datetime.datetime.now().isoformat())

        # Comment body paragraph
        cp = etree.SubElement(comment_elem, f"{{{W_NS}}}p")
        cr = etree.SubElement(cp, f"{{{W_NS}}}r")
        ct = etree.SubElement(cr, f"{{{W_NS}}}t")
        ct.text = comment_text

        # Insert commentRangeStart at beginning of paragraph
        range_start = etree.Element(f"{{{W_NS}}}commentRangeStart")
        range_start.set(f"{{{W_NS}}}id", comment_id)
        para_element.insert(0, range_start)

        # Insert commentRangeEnd + commentReference at end
        range_end = etree.SubElement(para_element, f"{{{W_NS}}}commentRangeEnd")
        range_end.set(f"{{{W_NS}}}id", comment_id)

        ref_run = etree.SubElement(para_element, f"{{{W_NS}}}r")
        ref_rpr = etree.SubElement(ref_run, f"{{{W_NS}}}rPr")
        ref_style = etree.SubElement(ref_rpr, f"{{{W_NS}}}rStyle")
        ref_style.set(f"{{{W_NS}}}val", "CommentReference")
        comment_ref = etree.SubElement(ref_run, f"{{{W_NS}}}commentReference")
        comment_ref.set(f"{{{W_NS}}}id", comment_id)

        # Update the comments part content
        self._comments_part._blob = etree.tostring(
            self._comments_element, xml_declaration=True, encoding="UTF-8"
        )

        return int(comment_id)


def _highlight_paragraph(para, color: str = "yellow"):
    """Add highlight to all runs in a paragraph."""
    from docx.oxml.ns import qn
    for run in para.findall(qn("w:r")):
        rpr = run.find(qn("w:rPr"))
        if rpr is None:
            rpr = etree.SubElement(run, qn("w:rPr"))
            run.insert(0, rpr)
        highlight = rpr.find(qn("w:highlight"))
        if highlight is None:
            highlight = etree.SubElement(rpr, qn("w:highlight"))
        highlight.set(qn("w:val"), color)


def _build_para_review_map(review_items: list[dict]) -> dict[int, list[tuple[dict, str]]]:
    """建立全局段落索引 → (review_item, per_para_reason) 的映射。

    兼容新旧两种格式：
    - 新格式: tender_locations[].global_para_indices + per_para_reasons
    - 旧格式: tender_locations[].para_indices
    """
    para_map: dict[int, list[tuple[dict, str]]] = {}
    for item in review_items:
        if item["result"] not in ("fail", "warning"):
            continue
        seen_paras: set[int] = set()
        for loc in item.get("tender_locations", []):
            indices = loc.get("global_para_indices", [])
            if not indices:
                indices = loc.get("para_indices", [])
            per_para_reasons = loc.get("per_para_reasons", {})
            for pi in indices:
                if pi not in seen_paras:
                    seen_paras.add(pi)
                    reason = per_para_reasons.get(pi, "") or per_para_reasons.get(str(pi), "")
                    para_map.setdefault(pi, []).append((item, reason))
    return para_map


def generate_review_docx(
    tender_file_path: str,
    review_items: list[dict],
    summary: dict,
    bid_filename: str = "",
    output_dir: str | None = None,
) -> str:
    """Generate annotated review docx with summary table + highlights + comments.

    Returns path to the generated docx file.
    """
    # Open tender file
    doc = Document(tender_file_path)

    # Build para_index → review_items mapping
    para_review_map = _build_para_review_map(review_items)

    # Add comments
    comment_mgr = _CommentManager(doc)

    # Get all body-level elements (paragraphs + tables) to match parser indexing
    from docx.oxml.ns import qn
    body = doc.element.body

    para_idx = 0
    for element in body:
        tag = element.tag
        if tag == qn("w:p"):
            # 与 docx_parser 保持一致：跳过空段落
            runs_text = []
            for r in element.findall(qn("w:r")):
                t = r.find(qn("w:t"))
                if t is not None and t.text:
                    runs_text.append(t.text)
            text = "".join(runs_text).strip()
            if not text:
                continue  # 空段落不计入索引

            if para_idx in para_review_map:
                entries = para_review_map[para_idx]
                highlight_color = "red" if any(item["result"] == "fail" for item, _ in entries) else "yellow"
                _highlight_paragraph(element, highlight_color)
                for item, per_para_reason in entries:
                    severity_label = {"critical": "废标条款", "major": "资格/编制要求", "minor": "评标标准"}.get(
                        item["severity"], "审查项"
                    )
                    result_label = {"pass": "合规", "fail": "不合规", "warning": "需注意", "error": "错误"}.get(
                        item["result"], item["result"]
                    )
                    # 优先使用该段落的独立 reason，fallback 到整体 reason
                    display_reason = per_para_reason or item.get("reason", "")
                    comment_text = (
                        f"[{severity_label} #{item.get('clause_index', '')}] "
                        f"置信度: {item.get('confidence', 0)}%\n"
                        f"判定: {_result_symbol(item['result'])} {result_label}\n"
                        f"条款: {item.get('clause_text', '')}\n"
                        f"原因: {display_reason}"
                    )
                    comment_mgr.add_comment(element, comment_text)
            para_idx += 1
        elif tag == qn("w:tbl"):
            # 表格作为一个段落计数（与 docx_parser 一致）
            # 表格不支持高亮/批注，但仍然需要递增索引
            para_idx += 1

    # Insert summary at beginning: build in main doc, then move to front
    original_count = len(list(body))
    tender_basename = os.path.basename(tender_file_path)
    _add_summary_section(doc, review_items, summary, bid_filename, tender_basename)

    # Move newly added summary elements (appended at end) to position 0
    new_elements = list(body)[original_count:]
    for i, elem in enumerate(new_elements):
        body.remove(elem)
        body.insert(i, elem)

    # Save
    if output_dir is None:
        output_dir = os.path.dirname(tender_file_path)
    tender_basename = os.path.basename(tender_file_path)
    output_filename = f"审查报告_{tender_basename}"
    output_path = os.path.join(output_dir, output_filename)
    doc.save(output_path)
    return output_path
