"""Tests for docx annotation generator."""
import os
import tempfile
from docx import Document
from src.reviewer.docx_annotator import generate_review_docx


def _create_test_docx(path: str):
    """Create a minimal test docx file."""
    doc = Document()
    doc.add_paragraph("第一章 投标函")
    doc.add_paragraph("致采购人：我方承诺按照招标文件要求提交投标文件。")
    doc.add_paragraph("第二章 技术方案")
    doc.add_paragraph("本系统采用分布式架构，具有高可用性。")
    doc.save(path)


def test_generate_review_docx_creates_file():
    """generate_review_docx produces a valid docx file."""
    with tempfile.TemporaryDirectory() as tmpdir:
        tender_path = os.path.join(tmpdir, "投标文件.docx")
        _create_test_docx(tender_path)

        review_items = [
            {
                "id": 0,
                "clause_index": 0,
                "source_module": "module_e",
                "clause_text": "投标文件须密封",
                "result": "fail",
                "confidence": 92,
                "reason": "未找到密封说明",
                "severity": "critical",
                "tender_locations": [{"chapter": "", "para_indices": [1], "text_snippet": ""}],
            }
        ]
        summary = {"total": 1, "pass": 0, "fail": 1, "warning": 0, "critical_fails": 1}

        output_path = generate_review_docx(
            tender_path, review_items, summary,
            bid_filename="招标文件.docx",
            output_dir=tmpdir,
        )

        assert os.path.exists(output_path)
        assert output_path.endswith(".docx")
        # Verify it's a valid docx
        doc = Document(output_path)
        assert len(doc.paragraphs) > 0


def test_generate_review_docx_has_summary_table():
    """Generated docx includes a summary table at the beginning."""
    with tempfile.TemporaryDirectory() as tmpdir:
        tender_path = os.path.join(tmpdir, "test.docx")
        _create_test_docx(tender_path)

        review_items = [{
            "id": 0, "clause_index": 0, "source_module": "module_e", "clause_text": "条款1",
            "result": "pass", "confidence": 90, "reason": "符合要求",
            "severity": "critical", "tender_locations": [],
        }]
        summary = {"total": 1, "pass": 1, "fail": 0, "warning": 0, "critical_fails": 0}

        output_path = generate_review_docx(
            tender_path, review_items, summary,
            bid_filename="test.docx", output_dir=tmpdir,
        )
        doc = Document(output_path)
        # Should have at least 1 table (summary table)
        assert len(doc.tables) >= 1


def test_generate_review_docx_has_word_comments():
    """Generated docx contains Word native comments for fail/warning items."""
    from zipfile import ZipFile

    with tempfile.TemporaryDirectory() as tmpdir:
        tender_path = os.path.join(tmpdir, "投标文件.docx")
        _create_test_docx(tender_path)

        review_items = [
            {
                "id": 0, "clause_index": 0, "source_module": "module_e",
                "clause_text": "投标文件须密封", "result": "fail",
                "confidence": 92, "reason": "未找到密封说明",
                "severity": "critical",
                "tender_locations": [{"chapter": "", "para_indices": [1], "text_snippet": ""}],
            }
        ]
        summary = {"total": 1, "pass": 0, "fail": 1, "warning": 0, "critical_fails": 1}

        output_path = generate_review_docx(
            tender_path, review_items, summary,
            bid_filename="招标文件.docx", output_dir=tmpdir,
        )
        # Verify comments.xml exists in the docx zip
        with ZipFile(output_path) as z:
            assert "word/comments.xml" in z.namelist()
            comments_xml = z.read("word/comments.xml").decode("utf-8")
            assert "未找到密封说明" in comments_xml


class TestBuildParaReviewMap:
    def test_new_format_global_para_indices(self):
        """新格式 global_para_indices 正确映射。"""
        from src.reviewer.docx_annotator import _build_para_review_map

        items = [{
            "result": "fail", "severity": "critical",
            "tender_locations": [{
                "batch_id": "/ch1#0",
                "path": "/ch1",
                "global_para_indices": [5, 6, 7],
                "text_snippet": "test",
            }],
        }]
        para_map = _build_para_review_map(items)
        assert 5 in para_map
        assert 6 in para_map
        assert 7 in para_map
        assert len(para_map[5]) == 1

    def test_old_format_para_indices(self):
        """旧格式 para_indices 仍然兼容。"""
        from src.reviewer.docx_annotator import _build_para_review_map

        items = [{
            "result": "fail", "severity": "critical",
            "tender_locations": [{
                "chapter": "第一章",
                "para_indices": [10, 11],
                "text_snippet": "test",
            }],
        }]
        para_map = _build_para_review_map(items)
        assert 10 in para_map
        assert 11 in para_map

    def test_dedup_same_para(self):
        """同一段落同一 item 只出现一次。"""
        from src.reviewer.docx_annotator import _build_para_review_map

        item = {
            "result": "fail", "severity": "critical",
            "tender_locations": [
                {"global_para_indices": [5], "text_snippet": "a"},
                {"global_para_indices": [5], "text_snippet": "b"},
            ],
        }
        para_map = _build_para_review_map([item])
        # 同一 item 在同一 para 只记录一次
        assert len(para_map[5]) == 1
