"""Tests for image extraction from docx/PDF files."""
import os
import tempfile
import pytest
from docx import Document
from docx.shared import Inches

from src.reviewer.image_extractor import extract_images


@pytest.fixture
def docx_with_image(tmp_path):
    """Create a docx file with an embedded image for testing."""
    doc = Document()
    doc.add_paragraph("第一章 投标函")
    # Create a minimal 1x1 PNG for testing
    import struct, zlib
    def make_png():
        """Generate a minimal 1x1 red PNG."""
        raw = b"\x00\xff\x00\x00"  # filter + RGB
        compressed = zlib.compress(raw)
        def chunk(ctype, data):
            c = ctype + data
            return struct.pack(">I", len(data)) + c + struct.pack(">I", zlib.crc32(c) & 0xFFFFFFFF)
        return (
            b"\x89PNG\r\n\x1a\n"
            + chunk(b"IHDR", struct.pack(">IIBBBBB", 1, 1, 8, 2, 0, 0, 0))
            + chunk(b"IDAT", compressed)
            + chunk(b"IEND", b"")
        )
    png_path = tmp_path / "test.png"
    png_path.write_bytes(make_png())
    doc.add_picture(str(png_path), width=Inches(2))
    doc.add_paragraph("资质证书如上图所示")
    doc.add_paragraph("第二章 技术方案")
    docx_path = tmp_path / "test_with_image.docx"
    doc.save(str(docx_path))
    return str(docx_path)


def test_extract_images_from_docx(docx_with_image, tmp_path):
    """Images are extracted from docx and saved to output dir."""
    output_dir = str(tmp_path / "images")
    images = extract_images(docx_with_image, output_dir)
    assert len(images) >= 1
    assert images[0]["filename"].endswith(".png") or images[0]["filename"].endswith(".jpeg")
    assert os.path.exists(images[0]["path"])
    # 独立图片段落的图片应关联到前一个文字段落 "第一章 投标函" (P0)，
    # 这样证书图片的描述会嵌入到正确的章节中，而非错误地关联到后续段落。
    assert images[0]["near_para_index"] == 0


def test_extract_images_no_images(tmp_path):
    """Docx without images returns empty list."""
    doc = Document()
    doc.add_paragraph("纯文本段落")
    docx_path = str(tmp_path / "no_image.docx")
    doc.save(docx_path)
    output_dir = str(tmp_path / "images")
    images = extract_images(docx_path, output_dir)
    assert images == []


def test_extract_images_from_pdf(tmp_path):
    """PDF image extraction returns a list (may be empty without pymupdf)."""
    # This is a graceful-degradation test: if pymupdf is not installed,
    # the function should return [] without raising.
    pdf_path = str(tmp_path / "test.pdf")
    # Create a minimal PDF (no images)
    with open(pdf_path, "wb") as f:
        f.write(b"%PDF-1.4\n1 0 obj<</Type/Catalog/Pages 2 0 R>>endobj\n"
                b"2 0 obj<</Type/Pages/Kids[3 0 R]/Count 1>>endobj\n"
                b"3 0 obj<</Type/Page/MediaBox[0 0 612 792]/Parent 2 0 R>>endobj\n"
                b"xref\n0 4\n0000000000 65535 f \n"
                b"0000000009 00000 n \n0000000058 00000 n \n0000000115 00000 n \n"
                b"trailer<</Size 4/Root 1 0 R>>\nstartxref\n190\n%%EOF")
    output_dir = str(tmp_path / "images")
    images = extract_images(pdf_path, output_dir)
    assert isinstance(images, list)
