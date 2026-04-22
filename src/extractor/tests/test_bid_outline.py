"""测试 bid_outline 四层流水线：合并/合成/绑定/编号/渲染/主入口。"""
import io
from unittest.mock import patch

from docx import Document
from docx.shared import RGBColor

from src.extractor.bid_outline import (
    _merge_skeleton_batches,
    _compose_outline_tree,
    _bind_sample_content,
    _normalize_title,
    _edit_distance_le2,
    _assign_numbering,
    _cn_numeral,
    _render_docx,
)


# ========== Layer 2 合并 ==========

def test_merge_skeleton_batches_concatenates_lists():
    b1 = {
        "composition_clause": {"found": True,
                                "items": [{"order": 1, "title": "商务文件"}]},
        "scoring_factors": [{"category": "技术", "title": "能力A", "sub_items": []}],
        "material_enumerations": [{"parent": "资质", "items": ["A"]}],
        "format_templates": [{"title": "投标函", "has_sample": True}],
        "dynamic_nodes": [],
    }
    b2 = {
        "composition_clause": {"found": False, "items": []},
        "scoring_factors": [{"category": "商务", "title": "能力B", "sub_items": []}],
        "material_enumerations": [{"parent": "证明", "items": ["B"]}],
        "format_templates": [{"title": "授权书", "has_sample": False}],
        "dynamic_nodes": [{"anchor": "业绩", "expansion_type": "customer_list",
                            "expansion_hint": "按客户展开"}],
    }
    merged = _merge_skeleton_batches([b1, b2])
    assert merged["composition_clause"]["found"] is True
    assert merged["composition_clause"]["items"][0]["title"] == "商务文件"
    assert len(merged["scoring_factors"]) == 2
    assert len(merged["material_enumerations"]) == 2
    assert len(merged["format_templates"]) == 2
    assert len(merged["dynamic_nodes"]) == 1


def test_merge_skeleton_batches_all_empty():
    merged = _merge_skeleton_batches([])
    assert merged["composition_clause"]["found"] is False
    assert merged["composition_clause"]["items"] == []
    assert merged["scoring_factors"] == []
    assert merged["material_enumerations"] == []
    assert merged["format_templates"] == []
    assert merged["dynamic_nodes"] == []


def test_merge_skeleton_batches_ignores_none_entries():
    b1 = {
        "composition_clause": {"found": True, "items": [{"order": 1, "title": "X"}]},
        "scoring_factors": [], "material_enumerations": [],
        "format_templates": [], "dynamic_nodes": [],
    }
    merged = _merge_skeleton_batches([None, b1, None])
    assert merged["composition_clause"]["found"] is True


def test_merge_skeleton_batches_preserves_first_composition_clause_wins():
    """多个批次都说 found=true 时保留第一个批次的 items（幂等，不累加）。"""
    b1 = {
        "composition_clause": {"found": True, "items": [{"order": 1, "title": "A"}]},
        "scoring_factors": [], "material_enumerations": [],
        "format_templates": [], "dynamic_nodes": [],
    }
    b2 = {
        "composition_clause": {"found": True, "items": [{"order": 1, "title": "B"}]},
        "scoring_factors": [], "material_enumerations": [],
        "format_templates": [], "dynamic_nodes": [],
    }
    merged = _merge_skeleton_batches([b1, b2])
    assert merged["composition_clause"]["found"] is True
    assert len(merged["composition_clause"]["items"]) == 1
    assert merged["composition_clause"]["items"][0]["title"] == "A"


# ========== Layer 3 合成 ==========

def test_compose_outline_tree_happy_path():
    layer1 = {"has_any_template": True, "templates": [
        {"title": "投标函", "type": "text", "content": "..."}
    ]}
    layer2 = {
        "composition_clause": {"found": True,
                                "items": [{"order": 1, "title": "投标函"}]},
        "scoring_factors": [], "material_enumerations": [],
        "format_templates": [], "dynamic_nodes": [],
    }
    fake_tree = {
        "title": "投标文件",
        "nodes": [{"title": "投标函", "level": 1,
                   "source": "format_template",
                   "has_sample": True, "dynamic": False,
                   "dynamic_hint": None, "children": []}]
    }
    with patch("src.extractor.bid_outline.call_qwen", return_value=fake_tree):
        result = _compose_outline_tree(layer1, layer2, settings=None)
    assert result == fake_tree


def test_compose_outline_tree_llm_returns_none():
    with patch("src.extractor.bid_outline.call_qwen", return_value=None):
        result = _compose_outline_tree(
            {"has_any_template": False, "templates": []},
            {"composition_clause": {"found": False, "items": []},
             "scoring_factors": [], "material_enumerations": [],
             "format_templates": [], "dynamic_nodes": []},
            settings=None,
        )
    assert result is None


def test_compose_outline_tree_llm_returns_malformed():
    with patch("src.extractor.bid_outline.call_qwen", return_value="not a dict"):
        assert _compose_outline_tree({}, {}, None) is None
    with patch("src.extractor.bid_outline.call_qwen", return_value={"foo": 1}):
        assert _compose_outline_tree({}, {}, None) is None


def test_compose_outline_tree_defaults_missing_title():
    """LLM 返回 nodes 但没有 title 字段时，主函数兜底填入'投标文件'。"""
    tree_no_title = {"nodes": []}
    with patch("src.extractor.bid_outline.call_qwen", return_value=tree_no_title):
        result = _compose_outline_tree({}, {}, None)
    assert result["title"] == "投标文件"
    assert result["nodes"] == []


def test_compose_outline_tree_passes_template_titles():
    """验证 Layer 1 templates 的 title 被提取并传入 user 文本。"""
    layer1 = {"has_any_template": True, "templates": [
        {"title": "投标函", "type": "text", "content": "X"},
        {"title": "开标一览表", "type": "standard_table",
         "columns": [], "rows": []},
    ]}
    captured = {}

    def _spy(messages, settings):
        captured["user"] = messages[-1]["content"]
        return {"title": "投标文件", "nodes": []}

    with patch("src.extractor.bid_outline.call_qwen", side_effect=_spy):
        _compose_outline_tree(layer1, {"composition_clause": {"found": False, "items": []},
                                        "scoring_factors": [], "material_enumerations": [],
                                        "format_templates": [], "dynamic_nodes": []}, None)
    assert "投标函" in captured["user"]
    assert "开标一览表" in captured["user"]


# ========== 样例绑定（模糊匹配） ==========

def test_normalize_title_strips_suffix_words_and_spaces():
    assert _normalize_title("投标函 格式") == "投标函"
    assert _normalize_title("  投标函模板  ") == "投标函"
    assert _normalize_title("开标一览表") == "开标一览"
    assert _normalize_title("授权委托书样表") == "授权委托书"
    assert _normalize_title("投标函样表格式") == "投标函"


def test_normalize_title_handles_empty_and_none():
    assert _normalize_title("") == ""
    assert _normalize_title(None) == ""


def test_edit_distance_le2():
    assert _edit_distance_le2("abc", "abc") is True  # 0
    assert _edit_distance_le2("abc", "abd") is True  # 1
    assert _edit_distance_le2("abc", "xyz") is False  # 3
    assert _edit_distance_le2("abcde", "ab") is False  # 3
    assert _edit_distance_le2("投标函", "投标书") is True  # 1
    assert _edit_distance_le2("投标函", "完全不同") is False


def test_bind_sample_content_exact_normalized_match():
    tree = {"nodes": [
        {"title": "投标函格式", "level": 1, "has_sample": True, "children": []}
    ]}
    layer1 = {"templates": [
        {"title": "投标函", "type": "text", "content": "TEMPLATE_CONTENT"}
    ]}
    _bind_sample_content(tree, layer1)
    assert tree["nodes"][0]["sample_content"] == {
        "type": "text", "content": "TEMPLATE_CONTENT"
    }


def test_bind_sample_content_substring_match():
    tree = {"nodes": [
        {"title": "开标报价一览表", "level": 1, "has_sample": True, "children": []}
    ]}
    layer1 = {"templates": [
        {"title": "开标一览表", "type": "standard_table",
         "columns": ["A"], "rows": [["1"]]}
    ]}
    _bind_sample_content(tree, layer1)
    bound = tree["nodes"][0]["sample_content"]
    assert bound["type"] == "standard_table"
    assert bound["columns"] == ["A"]
    assert bound["rows"] == [["1"]]


def test_bind_sample_content_no_match_leaves_none_and_keeps_has_sample():
    tree = {"nodes": [
        {"title": "完全不相关的标题XYZ", "level": 1, "has_sample": True, "children": []}
    ]}
    layer1 = {"templates": [
        {"title": "投标函", "type": "text", "content": "T"}
    ]}
    _bind_sample_content(tree, layer1)
    assert tree["nodes"][0]["sample_content"] is None
    assert tree["nodes"][0]["has_sample"] is True


def test_bind_sample_content_recurses_into_children():
    tree = {"nodes": [
        {"title": "附件", "level": 1, "has_sample": False, "children": [
            {"title": "投标函", "level": 2, "has_sample": True, "children": []}
        ]}
    ]}
    layer1 = {"templates": [
        {"title": "投标函", "type": "text", "content": "T"}
    ]}
    _bind_sample_content(tree, layer1)
    assert tree["nodes"][0]["children"][0]["sample_content"] is not None


def test_bind_sample_content_ignores_nodes_without_has_sample():
    """has_sample=False 的节点不应被污染 sample_content 字段。"""
    tree = {"nodes": [
        {"title": "投标函", "level": 1, "has_sample": False, "children": []}
    ]}
    layer1 = {"templates": [
        {"title": "投标函", "type": "text", "content": "T"}
    ]}
    _bind_sample_content(tree, layer1)
    assert "sample_content" not in tree["nodes"][0]


def test_bind_sample_content_handles_missing_layer1():
    """Layer 1 None / 空 templates 时不崩，has_sample=true 留 None。"""
    tree = {"nodes": [
        {"title": "投标函", "level": 1, "has_sample": True, "children": []}
    ]}
    _bind_sample_content(tree, None)
    assert tree["nodes"][0]["sample_content"] is None
    _bind_sample_content(tree, {"templates": []})
    assert tree["nodes"][0]["sample_content"] is None


# ========== 编号 ==========

def test_cn_numeral_1_to_20():
    expected = ["一","二","三","四","五","六","七","八","九","十",
                "十一","十二","十三","十四","十五","十六","十七","十八","十九","二十"]
    for i, e in enumerate(expected, start=1):
        assert _cn_numeral(i) == e


def test_cn_numeral_over_20_falls_back_to_arabic():
    assert _cn_numeral(21) == "21"
    assert _cn_numeral(99) == "99"


def test_cn_numeral_zero_or_negative_returns_arabic():
    """极端防御：非 1-20 范围统一降级。"""
    assert _cn_numeral(0) == "0"
    assert _cn_numeral(-1) == "-1"


def test_assign_numbering_three_level_tree():
    tree = {"nodes": [
        {"title": "附件", "level": 1, "children": [
            {"title": "资质证书", "level": 2, "children": [
                {"title": "A证书", "level": 3, "children": []},
                {"title": "B证书", "level": 3, "children": []},
            ]},
            {"title": "业绩", "level": 2, "children": []},
        ]},
        {"title": "技术部分", "level": 1, "children": [
            {"title": "质量方案", "level": 2, "children": []},
        ]},
    ]}
    _assign_numbering(tree)
    assert tree["nodes"][0]["number"] == "一、"
    assert tree["nodes"][0]["children"][0]["number"] == "1.1"
    assert tree["nodes"][0]["children"][0]["children"][0]["number"] == "1.1.1"
    assert tree["nodes"][0]["children"][0]["children"][1]["number"] == "1.1.2"
    assert tree["nodes"][0]["children"][1]["number"] == "1.2"
    assert tree["nodes"][1]["number"] == "二、"
    assert tree["nodes"][1]["children"][0]["number"] == "2.1"


def test_assign_numbering_empty_tree():
    tree = {"nodes": []}
    _assign_numbering(tree)  # 不崩溃
    assert tree["nodes"] == []


def test_assign_numbering_overflow_level1():
    """超过 20 个 level 1 节点时 21 起用阿拉伯数字。"""
    tree = {"nodes": [
        {"title": f"章节{i}", "level": 1, "children": []} for i in range(1, 22)
    ]}
    _assign_numbering(tree)
    assert tree["nodes"][19]["number"] == "二十、"
    assert tree["nodes"][20]["number"] == "21、"


# ========== Layer 4 渲染 ==========

def _build_rendered(tree: dict) -> Document:
    buf = io.BytesIO()
    _render_docx(tree, buf)
    buf.seek(0)
    return Document(buf)


def test_render_docx_writes_hierarchy_with_numbers():
    tree = {"title": "投标文件", "nodes": [
        {"number": "一、", "title": "附件", "level": 1,
         "has_sample": False, "dynamic": False, "children": [
             {"number": "1.1", "title": "资质", "level": 2,
              "has_sample": False, "dynamic": False, "children": [
                  {"number": "1.1.1", "title": "A证书", "level": 3,
                   "has_sample": False, "dynamic": False, "children": []}
              ]}
         ]}
    ]}
    doc = _build_rendered(tree)
    texts = [p.text for p in doc.paragraphs]
    assert "一、 附件" in texts
    assert "1.1 资质" in texts
    assert "1.1.1 A证书" in texts


def test_render_docx_embeds_text_sample():
    tree = {"title": "投标文件", "nodes": [
        {"number": "一、", "title": "投标函", "level": 1,
         "has_sample": True, "dynamic": False,
         "sample_content": {"type": "text",
                            "content": "致：[采购人]\n我方..."},
         "children": []}
    ]}
    doc = _build_rendered(tree)
    combined = "\n".join(p.text for p in doc.paragraphs)
    assert "致：[采购人]" in combined
    assert "我方..." in combined


def test_render_docx_embeds_table_sample():
    tree = {"title": "投标文件", "nodes": [
        {"number": "一、", "title": "开标一览表", "level": 1,
         "has_sample": True, "dynamic": False,
         "sample_content": {
             "type": "standard_table",
             "columns": ["项目", "金额"],
             "rows": [["A", "100"]],
         },
         "children": []}
    ]}
    doc = _build_rendered(tree)
    assert len(doc.tables) == 1
    tbl = doc.tables[0]
    assert tbl.rows[0].cells[0].text == "项目"
    assert tbl.rows[1].cells[1].text == "100"


def test_render_docx_dynamic_node_inserts_red_italic_hint():
    tree = {"title": "投标文件", "nodes": [
        {"number": "九、", "title": "类似项目业绩一览表", "level": 1,
         "has_sample": False, "dynamic": True,
         "dynamic_hint": "按客户/项目逐项展开",
         "children": []}
    ]}
    doc = _build_rendered(tree)
    hint_paras = [p for p in doc.paragraphs if "⚠" in p.text]
    assert len(hint_paras) == 1
    assert "按客户/项目逐项展开" in hint_paras[0].text
    run = hint_paras[0].runs[0]
    assert run.italic is True
    assert run.font.color.rgb == RGBColor(0xFF, 0x00, 0x00)


def test_render_docx_dynamic_without_hint_falls_back():
    """dynamic=True 但无 dynamic_hint 时应有默认提示。"""
    tree = {"title": "投标文件", "nodes": [
        {"number": "一、", "title": "X", "level": 1,
         "has_sample": False, "dynamic": True, "children": []}
    ]}
    doc = _build_rendered(tree)
    hint_paras = [p for p in doc.paragraphs if "⚠" in p.text]
    assert len(hint_paras) == 1


def test_render_docx_accepts_file_path(tmp_path):
    tree = {"title": "投标文件", "nodes": [
        {"number": "一、", "title": "X", "level": 1,
         "has_sample": False, "dynamic": False, "children": []}
    ]}
    out = tmp_path / "o.docx"
    _render_docx(tree, str(out))
    assert out.exists() and out.stat().st_size > 0


def test_render_docx_empty_nodes_writes_only_title():
    tree = {"title": "投标文件", "nodes": []}
    doc = _build_rendered(tree)
    texts = [p.text for p in doc.paragraphs]
    assert "投标文件" in texts
