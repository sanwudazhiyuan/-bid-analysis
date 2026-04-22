"""bid_outline: 投标文件大纲生成模块（四层流水线）

Layer 1: 格式样例抽取（沿用 config/prompts/bid_format.txt，复用 bid_format._first_pass）
Layer 2: 骨架信号抽取（本模块新增）
Layer 3: 目录合成（单次 LLM 调用）
Layer 4: docx 渲染（纯代码，无 LLM）

对外入口：`extract_bid_outline(tagged_paragraphs, settings, embeddings_map,
module_embedding, modules_context=None) -> dict | None`。
"""
import json as _json
import logging
import re as _re
from pathlib import Path

from src.models import TaggedParagraph
from src.extractor.base import (
    load_prompt_template,
    build_messages,
    build_input_text,
    call_qwen,
    batch_by_count,
)
from src.extractor.scoring import filter_paragraphs_by_score

logger = logging.getLogger(__name__)

_CONFIG_DIR = Path(__file__).parent.parent.parent / "config"
SKELETON_PROMPT_PATH = _CONFIG_DIR / "prompts" / "bid_format_skeleton.txt"
COMPOSE_PROMPT_PATH = _CONFIG_DIR / "prompts" / "bid_format_compose.txt"

BATCH_SIZE = 40
TOKEN_SAFETY_CAP = 100_000
MIN_FILTER_COUNT = 50


def _empty_skeleton() -> dict:
    """返回一份新的空骨架结构（每次返回独立 dict 避免共享）。"""
    return {
        "composition_clause": {"found": False, "items": []},
        "scoring_factors": [],
        "material_enumerations": [],
        "format_templates": [],
        "dynamic_nodes": [],
    }


def _merge_skeleton_batches(batch_results: list) -> dict:
    """合并多批 Layer 2 输出为单份 JSON。

    - composition_clause：首个 found=True 的批次获胜，保留其 items（通常招标方只写一处）
    - 其余 4 类：列表拼接，**不做去重**
    - 非 dict / None 的条目跳过
    """
    merged = _empty_skeleton()
    for b in batch_results:
        if not isinstance(b, dict):
            continue
        cc = b.get("composition_clause") or {}
        if cc.get("found") and not merged["composition_clause"]["found"]:
            merged["composition_clause"] = {
                "found": True,
                "items": list(cc.get("items") or []),
            }
        for k in ("scoring_factors", "material_enumerations",
                  "format_templates", "dynamic_nodes"):
            v = b.get(k)
            if isinstance(v, list):
                merged[k].extend(v)
    return merged


def _extract_skeleton_signals(
    tagged_paragraphs: list[TaggedParagraph],
    settings: dict | None = None,
    embeddings_map: dict[int, list[float]] | None = None,
    module_embedding: list[float] | None = None,
) -> dict | None:
    """Layer 2：关键词过滤 → 按段落数分批 → LLM 抽 5 类信号 → 合并。

    返回：
    - 正常：合并后的骨架 dict（5 个字段都存在）
    - 过滤后段落为空：返回空骨架（不视为失败，交主入口做三重空判定）
    - 全部批次 LLM 都失败：返回 None
    """
    filtered, score_map = filter_paragraphs_by_score(
        tagged_paragraphs, "bid_format_skeleton",
        embeddings_map=embeddings_map,
        module_embedding=module_embedding,
        min_count=MIN_FILTER_COUNT,
    )
    if not filtered:
        logger.warning("bid_outline.layer2: 未筛选到相关段落")
        return _empty_skeleton()

    logger.info("bid_outline.layer2: 筛选到 %d 个相关段落 (共 %d)",
                len(filtered), len(tagged_paragraphs))

    system_prompt = load_prompt_template(str(SKELETON_PROMPT_PATH))
    batches = batch_by_count(filtered, batch_size=BATCH_SIZE,
                             token_safety_cap=TOKEN_SAFETY_CAP)

    batch_results: list = []
    for i, batch in enumerate(batches):
        batch_text = build_input_text(batch, score_map)
        messages = build_messages(system=system_prompt, user=batch_text)
        logger.debug("bid_outline.layer2: 调用第 %d/%d 批 (段落数=%d)",
                     i + 1, len(batches), len(batch))
        result = call_qwen(messages, settings)
        batch_results.append(result if isinstance(result, dict) else None)

    if all(r is None for r in batch_results):
        logger.error("bid_outline.layer2: 所有批次 LLM 返回均失败")
        return None

    return _merge_skeleton_batches(batch_results)


# ========== Layer 3：目录合成 ==========

def _compose_outline_tree(
    layer1_result: dict | None,
    layer2_result: dict | None,
    settings: dict | None,
) -> dict | None:
    """Layer 3：输入 Layer 1 样例 title 列表 + Layer 2 结构信号，输出多级目录树。

    输出未编号、未绑定 sample_content。LLM 返回非 dict 或缺 `nodes` 返回 None。
    """
    template_titles: list[str] = []
    if isinstance(layer1_result, dict) and layer1_result.get("has_any_template"):
        for t in layer1_result.get("templates") or []:
            if isinstance(t, dict) and t.get("title"):
                template_titles.append(t["title"])

    safe_layer2 = layer2_result if isinstance(layer2_result, dict) else _empty_skeleton()

    payload = {
        "layer1_template_titles": template_titles,
        "layer2_skeleton": safe_layer2,
    }

    system_prompt = load_prompt_template(str(COMPOSE_PROMPT_PATH))
    user_text = _json.dumps(payload, ensure_ascii=False, indent=2)
    messages = build_messages(system=system_prompt, user=user_text)

    result = call_qwen(messages, settings)
    if not isinstance(result, dict) or "nodes" not in result:
        logger.error("bid_outline.layer3: LLM 返回非法结构: %s", type(result).__name__)
        return None
    if "title" not in result:
        result["title"] = "投标文件"
    return result


# ========== 样例内容绑定（模糊匹配） ==========

# 尾部后缀词（按长度降序排列，确保最长匹配优先，如"样表" 先于 "表"）
_TITLE_SUFFIX_WORDS = ("格式", "模板", "样表", "样式", "表")


def _normalize_title(title) -> str:
    """归一化标题：去空格、剥离尾部后缀词（循环直至稳定）。"""
    t = (title or "")
    if not isinstance(t, str):
        return ""
    t = _re.sub(r"\s+", "", t.strip())
    changed = True
    while changed:
        changed = False
        for sw in _TITLE_SUFFIX_WORDS:
            if t.endswith(sw) and len(t) > len(sw):
                t = t[: -len(sw)]
                changed = True
                break  # 重新从最长后缀词开始检查
    return t


def _edit_distance_le2(a: str, b: str) -> bool:
    """判定编辑距离 ≤ 2；不追求最优、朴素 DP。"""
    if a == b:
        return True
    la, lb = len(a), len(b)
    if abs(la - lb) > 2:
        return False
    prev = list(range(lb + 1))
    for i in range(1, la + 1):
        cur = [i] + [0] * lb
        for j in range(1, lb + 1):
            cost = 0 if a[i - 1] == b[j - 1] else 1
            cur[j] = min(cur[j - 1] + 1, prev[j] + 1, prev[j - 1] + cost)
        prev = cur
    return prev[lb] <= 2


def _find_template_for_title(node_title: str, templates: list) -> dict | None:
    """按归一化精确 → 子串 → 编辑距离 ≤ 2 顺序查找匹配模板。"""
    if not templates:
        return None
    target = _normalize_title(node_title)
    if not target:
        return None
    # 1. 归一化精确匹配
    for t in templates:
        if _normalize_title(t.get("title", "")) == target:
            return t
    # 2. 子串（双向）
    for t in templates:
        tnorm = _normalize_title(t.get("title", ""))
        if tnorm and (tnorm in target or target in tnorm):
            return t
    # 3. 编辑距离
    for t in templates:
        tnorm = _normalize_title(t.get("title", ""))
        if tnorm and _edit_distance_le2(tnorm, target):
            return t
    return None


def _extract_sample_payload(template: dict) -> dict:
    """从 Layer 1 template 剥出 Layer 4 渲染所需的 sample_content。"""
    t = template.get("type", "text")
    if t == "standard_table":
        return {
            "type": "standard_table",
            "columns": list(template.get("columns") or []),
            "rows": [list(r) for r in (template.get("rows") or [])],
        }
    return {"type": "text", "content": template.get("content", "") or ""}


def _bind_sample_content(tree: dict, layer1_result: dict | None) -> None:
    """递归为 tree 中 has_sample=true 的节点绑定 sample_content（原地修改）。

    - has_sample=true 但无匹配 → sample_content=None，has_sample 标记保留
    - has_sample=false → 不触碰节点字段
    """
    templates: list = []
    if isinstance(layer1_result, dict):
        templates = [t for t in (layer1_result.get("templates") or [])
                     if isinstance(t, dict)]

    def _walk(node: dict) -> None:
        if node.get("has_sample"):
            match = _find_template_for_title(node.get("title", ""), templates)
            if match:
                node["sample_content"] = _extract_sample_payload(match)
            else:
                node["sample_content"] = None
                logger.warning(
                    "bid_outline.bind: has_sample=True 但未匹配到样例 title=%s",
                    node.get("title"),
                )
        for child in node.get("children") or []:
            _walk(child)

    for n in tree.get("nodes") or []:
        _walk(n)


# ========== 编号后处理 ==========

_CN_NUMERALS = [
    "一", "二", "三", "四", "五", "六", "七", "八", "九", "十",
    "十一", "十二", "十三", "十四", "十五", "十六", "十七", "十八", "十九", "二十",
]


def _cn_numeral(n: int) -> str:
    """1~20 返回中文数字，其它降级为阿拉伯数字字符串。"""
    if isinstance(n, int) and 1 <= n <= 20:
        return _CN_NUMERALS[n - 1]
    return str(n)


def _assign_numbering(tree: dict) -> None:
    """深度优先为 tree.nodes 每个节点赋 number 字段（原地修改）。

    - Level 1：`{中文数字}、`
    - Level 2+：`{父阿拉伯编号}.{本级序号}`
    """
    for i, node in enumerate(tree.get("nodes") or [], start=1):
        node["number"] = f"{_cn_numeral(i)}、"
        _assign_sub(node, prefix=str(i))


def _assign_sub(parent: dict, prefix: str) -> None:
    for j, child in enumerate(parent.get("children") or [], start=1):
        child["number"] = f"{prefix}.{j}"
        _assign_sub(child, prefix=f"{prefix}.{j}")


# ========== Layer 4：docx 渲染 ==========

from docx import Document as _DocxDocument  # noqa: E402
from docx.shared import RGBColor as _RGBColor  # noqa: E402


def _render_docx(tree: dict, output) -> None:
    """把目录树渲染为 docx。

    output: 路径（str/Path）或可写 file-like 对象。
    """
    doc = _DocxDocument()
    doc.add_heading(tree.get("title", "投标文件"), level=0)
    for node in tree.get("nodes") or []:
        _render_node(doc, node)
    doc.save(output)


def _render_node(doc, node: dict) -> None:
    level = min(max(int(node.get("level", 1) or 1), 1), 3)
    number = node.get("number", "")
    title = node.get("title", "")
    heading_text = f"{number} {title}".strip()
    doc.add_heading(heading_text, level=level)

    # 动态节点提示：红字 + 斜体
    if node.get("dynamic"):
        hint = node.get("dynamic_hint") or "按实际情况展开"
        sample_n1 = f'"{number}.1 示例一"' if number else '"示例一"'
        sample_n2 = f'"{number}.2 示例二"' if number else '"示例二"'
        line = f"[⚠ 此节需{hint}，例如 {sample_n1} / {sample_n2}]"
        p = doc.add_paragraph()
        run = p.add_run(line)
        run.italic = True
        run.font.color.rgb = _RGBColor(0xFF, 0x00, 0x00)

    # 样例嵌入
    sc = node.get("sample_content")
    if node.get("has_sample") and isinstance(sc, dict):
        if sc.get("type") == "standard_table":
            _add_table(doc, sc.get("columns") or [], sc.get("rows") or [])
        else:
            content = sc.get("content") or ""
            for line in content.split("\n"):
                doc.add_paragraph(line)

    children = node.get("children") or []
    if not children:
        doc.add_paragraph("")
    else:
        for child in children:
            _render_node(doc, child)


def _add_table(doc, columns: list, rows: list) -> None:
    if not columns:
        return
    ncols = len(columns)
    tbl = doc.add_table(rows=1 + len(rows), cols=ncols)
    tbl.style = "Table Grid"
    for i, c in enumerate(columns):
        tbl.rows[0].cells[i].text = str(c)
    for ri, row in enumerate(rows, start=1):
        padded = list(row) + [""] * (ncols - len(row))
        for ci, cell in enumerate(padded[:ncols]):
            tbl.rows[ri].cells[ci].text = str(cell)


# ========== 主入口：四层并行编排 ==========

from concurrent.futures import ThreadPoolExecutor  # noqa: E402


def _layer1_sections_to_templates(sections) -> list:
    """把旧 bid_format.txt 的 sections 结构展平为 templates 数组。

    sections 可能含 type=group 的分组节点（带 children），需要递归展平。
    """
    out: list = []
    if not sections:
        return out
    for s in sections:
        if not isinstance(s, dict):
            continue
        st = s.get("type")
        if st == "group":
            out.extend(_layer1_sections_to_templates(s.get("children", [])))
        elif st == "standard_table":
            out.append({
                "title": s.get("title", ""),
                "type": "standard_table",
                "columns": list(s.get("columns") or []),
                "rows": [list(r) for r in (s.get("rows") or [])],
            })
        else:
            out.append({
                "title": s.get("title", ""),
                "type": "text",
                "content": s.get("content", "") or "",
            })
    return out


def _run_layer1(
    tagged_paragraphs: list[TaggedParagraph],
    settings: dict | None,
    embeddings_map: dict[int, list[float]] | None,
    module_embedding: list[float] | None,
) -> dict:
    """Layer 1：复用现有 bid_format 的第一次 LLM 调用。

    返回统一 schema: {"has_any_template": bool, "templates": [...]}。
    遇到任何异常不抛、返回空结构，由主入口决定是否继续。
    """
    try:
        # 延迟导入避免循环依赖（bid_format.py 会被 Task 13 精简为 helper 模块）
        from src.extractor.bid_format import _filter_paragraphs, _first_pass

        filtered, score_map = _filter_paragraphs(
            tagged_paragraphs,
            embeddings_map=embeddings_map,
            module_embedding=module_embedding,
        )
        if not filtered:
            return {"has_any_template": False, "templates": []}
        raw = _first_pass(filtered, score_map, settings)
        if not isinstance(raw, dict):
            return {"has_any_template": False, "templates": []}
        if raw.get("has_template") is False:
            return {"has_any_template": False, "templates": []}

        templates = _layer1_sections_to_templates(raw.get("sections", []))
        return {
            "has_any_template": len(templates) > 0,
            "templates": templates,
        }
    except Exception as e:
        logger.warning("bid_outline.layer1 失败: %s", e)
        return {"has_any_template": False, "templates": []}


def _is_triple_empty(layer1: dict, layer2: dict) -> bool:
    """所有 5 类结构信号都为空时视为三重空。"""
    if not isinstance(layer1, dict) or not isinstance(layer2, dict):
        return True
    if layer1.get("templates"):
        return False
    cc = layer2.get("composition_clause") or {}
    if cc.get("found"):
        return False
    if layer2.get("scoring_factors"):
        return False
    if layer2.get("material_enumerations"):
        return False
    if layer2.get("format_templates"):
        return False
    if layer2.get("dynamic_nodes"):
        return False
    return True


def extract_bid_outline(
    tagged_paragraphs: list[TaggedParagraph],
    settings: dict | None = None,
    embeddings_map: dict[int, list[float]] | None = None,
    module_embedding: list[float] | None = None,
    modules_context: dict | None = None,  # 兼容旧签名，不再使用
) -> dict | None:
    """四层流水线主入口。返回目录树 JSON 或 None（失败/空信号）。

    docx 落盘不在此函数内完成，调用方需渲染时调用 ``_render_docx(tree, path)``。
    """
    # Layer 1 与 Layer 2 并行（共两个 LLM 任务）
    with ThreadPoolExecutor(max_workers=2) as ex:
        fut_l1 = ex.submit(
            _run_layer1, tagged_paragraphs, settings,
            embeddings_map, module_embedding,
        )
        fut_l2 = ex.submit(
            _extract_skeleton_signals, tagged_paragraphs, settings,
            embeddings_map, module_embedding,
        )
        try:
            layer1 = fut_l1.result()
        except Exception as e:
            logger.warning("bid_outline: Layer 1 抛异常: %s，视作无样例继续", e)
            layer1 = {"has_any_template": False, "templates": []}
        try:
            layer2 = fut_l2.result()
        except Exception as e:
            logger.error("bid_outline: Layer 2 抛异常: %s", e)
            layer2 = None

    if layer2 is None:
        logger.error("bid_outline: Layer 2 返回 None，整体失败")
        return None

    if _is_triple_empty(layer1, layer2):
        logger.error("bid_outline: 三重空信号，无法生成大纲")
        return None

    tree = _compose_outline_tree(layer1, layer2, settings)
    if tree is None:
        logger.error("bid_outline: Layer 3 返回 None")
        return None

    _bind_sample_content(tree, layer1)
    _assign_numbering(tree)

    return tree
