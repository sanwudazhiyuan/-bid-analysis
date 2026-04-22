"""投标文件大纲生成器：将提取的投标文件大纲渲染为 .docx。

支持两种数据结构：
- 旧结构：bid_format = {"sections": [...]}（type=group/table/text）
- 新结构：bid_format = {"title": ..., "nodes": [...]}（level/number/dynamic/sample_content）

优先按新结构渲染；若 nodes 不存在则降级到旧 sections 逻辑。
"""
import os

from docx import Document
from docx.shared import RGBColor

from src.generator.style_manager import StyleManager
from src.generator.table_builder import TableBuilder


# ---- 新结构渲染（nodes 树） ----

def _render_node(node: dict, doc: Document, style_mgr: StyleManager,
                  table_builder: TableBuilder) -> None:
    """递归渲染一个 outline node（新结构）。"""
    number = node.get("number", "")
    title = node.get("title", "")
    level = min(max(int(node.get("level", 1) or 1), 1), 3)

    # 标题
    heading_text = f"{number} {title}".strip() if number else title
    if heading_text:
        heading_para = doc.add_paragraph()
        if level == 1:
            run = heading_para.add_run(heading_text)
            style_mgr.apply_run_style(run, "heading2")
        elif level == 2:
            run = heading_para.add_run(heading_text)
            style_mgr.apply_run_style(run, "heading3")
        else:
            run = heading_para.add_run(heading_text)
            style_mgr.apply_run_style(run, "body")
            run.bold = True

    # 动态节点提示：红色斜体
    if node.get("dynamic"):
        hint = node.get("dynamic_hint") or "按实际情况展开"
        sample_n1 = f'"{number}.1 示例一"' if number else '"示例一"'
        sample_n2 = f'"{number}.2 示例二"' if number else '"示例二"'
        line = f"[此节需{hint}，例如 {sample_n1} / {sample_n2}]"
        p = doc.add_paragraph()
        r = p.add_run(line)
        r.italic = True
        r.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
        style_mgr.apply_run_style(r, "body")

    # 样例嵌入
    sc = node.get("sample_content")
    if node.get("has_sample") and isinstance(sc, dict):
        if sc.get("type") == "standard_table":
            table_builder.build(sc, doc)
        else:
            content = sc.get("content", "") or ""
            if content:
                for line in content.split("\n"):
                    para = doc.add_paragraph()
                    run = para.add_run(line)
                    style_mgr.apply_run_style(run, "body")

    # 子节点
    children = node.get("children") or []
    if not children:
        doc.add_paragraph()
    else:
        for child in children:
            _render_node(child, doc, style_mgr, table_builder)


def _render_node_toc(nodes: list[dict], doc: Document, style_mgr: StyleManager) -> None:
    """为新结构渲染文本目录。"""
    toc_title = doc.add_paragraph()
    toc_run = toc_title.add_run("目  录")
    style_mgr.apply_run_style(toc_run, "heading2")
    toc_title.alignment = 1

    def _walk(items: list[dict], depth: int = 0):
        for item in items:
            number = item.get("number", "")
            title = item.get("title", "")
            indent = "    " * depth
            if depth == 0:
                line = f"{indent}{number}{title}" if number else f"{indent}{title}"
            else:
                line = f"{indent}{number} {title}" if number else f"{indent}{title}"
            para = doc.add_paragraph()
            run = para.add_run(line)
            style_mgr.apply_run_style(run, "body")
            if depth == 0:
                run.bold = True
            _walk(item.get("children") or [], depth + 1)

    _walk(nodes)
    doc.add_paragraph()


# ---- 旧结构渲染（sections） ----

def _render_section(section: dict, doc: Document, style_mgr: StyleManager,
                    table_builder: TableBuilder, level: int = 1) -> None:
    """递归渲染一个 section 节点（旧结构）。"""
    number = section.get("number", "")
    title = section.get("title", "")
    section_type = section.get("type", "")

    if title:
        heading_text = f"{number}、{title}" if level == 1 and number else f"{number} {title}" if number else title
        heading_para = doc.add_paragraph()
        if level == 1:
            heading_run = heading_para.add_run(heading_text)
            style_mgr.apply_run_style(heading_run, "heading2")
        elif level == 2:
            heading_run = heading_para.add_run(heading_text)
            style_mgr.apply_run_style(heading_run, "heading3")
        else:
            heading_run = heading_para.add_run(heading_text)
            style_mgr.apply_run_style(heading_run, "body")
            heading_run.bold = True

    if section_type == "group":
        for child in section.get("children", []):
            _render_section(child, doc, style_mgr, table_builder, level + 1)
        return

    if section_type in ("key_value_table", "standard_table"):
        table_builder.build(section, doc)
    elif section_type in ("text", "template"):
        content = section.get("content", "")
        if content:
            para = doc.add_paragraph()
            run = para.add_run(content)
            style_mgr.apply_run_style(run, "body")


def _render_toc(sections: list[dict], doc: Document, style_mgr: StyleManager) -> None:
    """旧结构：在正文前渲染文本目录。"""
    toc_title = doc.add_paragraph()
    toc_run = toc_title.add_run("目  录")
    style_mgr.apply_run_style(toc_run, "heading2")
    toc_title.alignment = 1

    def _walk(items: list[dict], depth: int = 0):
        for item in items:
            number = item.get("number", "")
            title = item.get("title", "")
            indent = "    " * depth
            if depth == 0:
                line = f"{indent}{number}、{title}" if number else f"{indent}{title}"
            else:
                line = f"{indent}{number} {title}" if number else f"{indent}{title}"
            para = doc.add_paragraph()
            run = para.add_run(line)
            style_mgr.apply_run_style(run, "body")
            if depth == 0:
                run.bold = True
            if item.get("type") == "group":
                _walk(item.get("children", []), depth + 1)

    _walk(sections)
    doc.add_paragraph()


# ---- 主入口 ----

def render_format(data: dict, output_path: str) -> None:
    """将投标文件大纲渲染为 .docx。

    优先按新结构（nodes 树）渲染；若 nodes 不存在则降级到旧 sections。
    """
    parent = os.path.dirname(output_path)
    if parent:
        os.makedirs(parent, exist_ok=True)

    doc = Document()
    style_mgr = StyleManager()
    table_builder = TableBuilder(style_manager=style_mgr)

    modules = data.get("modules", {})
    bid_format = modules.get("bid_format")

    if bid_format is None:
        para = doc.add_paragraph()
        run = para.add_run("[投标文件大纲模块无数据]")
        run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
        doc.save(output_path)
        return

    if bid_format.get("status") == "failed":
        error = bid_format.get("error", "未知错误")
        para = doc.add_paragraph()
        run = para.add_run(f"[投标文件大纲提取失败: {error}]")
        run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
        doc.save(output_path)
        return

    # 标题
    tree_title = bid_format.get("title", "投标文件大纲")
    title_para = doc.add_paragraph()
    title_run = title_para.add_run(tree_title)
    style_mgr.apply_run_style(title_run, "heading1")
    title_para.alignment = 1

    # 新结构：nodes 树
    nodes = bid_format.get("nodes")
    if nodes:
        _render_node_toc(nodes, doc, style_mgr)
        for node in nodes:
            _render_node(node, doc, style_mgr, table_builder)
    else:
        # 降级到旧结构：sections
        sections = bid_format.get("sections", [])
        if sections:
            _render_toc(sections, doc, style_mgr)
        for section in sections:
            _render_section(section, doc, style_mgr, table_builder, level=1)

    doc.save(output_path)