"""资料清单生成器：将提取的资料清单渲染为 .docx。

输出包含分类材料表格（资格证明、技术能力等）。
"""

import os

from docx import Document
from docx.shared import RGBColor

from src.generator.style_manager import StyleManager
from src.generator.table_builder import TableBuilder


def render_checklist(data: dict, output_path: str) -> None:
    """将提取结果中的资料清单渲染为 .docx。

    Args:
        data: 包含 modules 的提取结果 dict
        output_path: 输出 .docx 文件路径
    """
    parent = os.path.dirname(output_path)
    if parent:
        os.makedirs(parent, exist_ok=True)

    doc = Document()
    style_mgr = StyleManager()
    table_builder = TableBuilder(style_manager=style_mgr)

    # 大标题
    title_para = doc.add_paragraph()
    title_run = title_para.add_run("资料清单")
    style_mgr.apply_run_style(title_run, "heading1")
    title_para.alignment = 1

    modules = data.get("modules", {})
    checklist = modules.get("checklist")

    if checklist is None:
        para = doc.add_paragraph()
        run = para.add_run("[资料清单模块无数据]")
        run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
        doc.save(output_path)
        return

    if checklist.get("status") == "failed":
        error = checklist.get("error", "未知错误")
        para = doc.add_paragraph()
        run = para.add_run(f"[资料清单提取失败: {error}]")
        run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
        doc.save(output_path)
        return

    sections = checklist.get("sections", [])
    for section in sections:
        section_type = section.get("type", "")
        section_title = section.get("title", "")

        if section_title:
            heading_para = doc.add_paragraph()
            heading_run = heading_para.add_run(section_title)
            style_mgr.apply_run_style(heading_run, "heading2")

        if section_type in ("key_value_table", "standard_table"):
            table_builder.build(section, doc)
        elif section_type in ("template", "text"):
            content = section.get("content", "")
            if content:
                para = doc.add_paragraph()
                run = para.add_run(content)
                style_mgr.apply_run_style(run, "body")

    doc.save(output_path)
