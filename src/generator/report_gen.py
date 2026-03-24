"""分析报告生成器：将校对后的 JSON 渲染为 .docx 分析报告。

输出包含 A-G 模块标题、编号子标题和表格。
不包含 bid_format 和 checklist 模块（它们有独立生成器）。
子标题下只输出表格，text 类型内容也转为表格渲染。
每个表格末尾添加"确认"勾选列。
"""

import os
import re

from docx import Document
from docx.shared import RGBColor, Pt

from src.generator.style_manager import StyleManager
from src.generator.table_builder import TableBuilder

# 分析报告中排除的模块（有独立生成器）
_EXCLUDED_MODULES = {"bid_format", "checklist"}


def _extract_module_letter(title: str) -> str:
    """从模块标题中提取字母前缀，如 'A. 项目概况' → 'A'"""
    m = re.match(r"([A-G])\.", title)
    return m.group(1) if m else ""


def _add_checkbox_column(section: dict) -> dict:
    """为 section 添加"确认"勾选列，返回新 dict。"""
    section = dict(section)
    columns = list(section.get("columns", []))
    rows = [list(row) for row in section.get("rows", [])]

    columns.append("确认")
    for row in rows:
        row.append("☐")

    section["columns"] = columns
    section["rows"] = rows
    return section


def render_report(data: dict, output_path: str) -> None:
    """将提取/校对结果渲染为分析报告 .docx。

    Args:
        data: 包含 modules 的提取结果 dict
        output_path: 输出 .docx 文件路径
    """
    parent = os.path.dirname(output_path)
    if parent:
        os.makedirs(parent, exist_ok=True)

    doc = Document()
    style_mgr = StyleManager()
    table_builder = TableBuilder(style_manager=style_mgr)

    # 报告大标题
    title_para = doc.add_paragraph()
    title_run = title_para.add_run("招标文件分析报告")
    style_mgr.apply_run_style(title_run, "heading1")
    title_para.alignment = 1  # 居中

    modules = data.get("modules", {})

    for module_key, module_data in modules.items():
        # 排除投标文件格式和资料清单
        if module_key in _EXCLUDED_MODULES:
            continue

        # None 模块
        if module_data is None:
            para = doc.add_paragraph()
            run = para.add_run(f"[{module_key}] 无数据（跳过）")
            run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
            continue

        # 失败模块
        if module_data.get("status") == "failed":
            error = module_data.get("error", "未知错误")
            para = doc.add_paragraph()
            run = para.add_run(
                f"[{module_key} 提取失败: {error}]"
            )
            run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
            continue

        # 正常模块 — 渲染标题
        title = module_data.get("title", module_key)
        heading_para = doc.add_paragraph()
        heading_run = heading_para.add_run(title)
        style_mgr.apply_run_style(heading_run, "heading2")

        # 提取模块字母前缀用于子标题编号
        letter = _extract_module_letter(title)

        # 渲染 sections
        sections = module_data.get("sections", [])
        _render_sections(doc, sections, table_builder, style_mgr, letter)

    doc.save(output_path)


def _render_sections(
    doc: Document,
    sections: list,
    table_builder: TableBuilder,
    style_mgr: StyleManager,
    module_letter: str,
) -> None:
    """渲染 sections 列表，使用 section 的 id 字段作为编号前缀。"""
    for idx, section in enumerate(sections, 1):
        section_type = section.get("type", "")
        section_title = section.get("title", "")
        section_id = section.get("id", "")

        # 子标题：优先使用 id 字段作为编号前缀
        if section_title:
            if section_id:
                numbered_title = f"{section_id} {section_title}"
            elif module_letter:
                numbered_title = f"{module_letter}.{idx} {section_title}"
            else:
                numbered_title = section_title
            sub_para = doc.add_paragraph()
            sub_run = sub_para.add_run(numbered_title)
            style_mgr.apply_run_style(sub_run, "heading3")

        # 渲染内容 — 只允许表格，并添加勾选列
        if section_type in ("key_value_table", "standard_table"):
            section_no_title = {k: v for k, v in section.items() if k != "title"}
            section_with_check = _add_checkbox_column(section_no_title)
            table_builder.build(section_with_check, doc)
        elif section_type in ("text", "template"):
            content = section.get("content", "")
            if content:
                _render_text_as_table(doc, section_title, content, table_builder, style_mgr)

        # note 字段渲染
        note = section.get("note", "")
        if note:
            _render_note(doc, note, style_mgr)

        # 递归处理子 sections（parent type）
        sub_sections = section.get("sections", [])
        if sub_sections:
            _render_sections(doc, sub_sections, table_builder, style_mgr, module_letter)


def _render_text_as_table(
    doc: Document,
    title: str,
    content: str,
    table_builder: TableBuilder,
    style_mgr: StyleManager,
) -> None:
    """将文本内容转为表格渲染（单列表格：标题+内容+勾选列）。"""
    section_as_table = {
        "type": "key_value_table",
        "columns": ["条款", "内容"],
        "rows": [[title or "详情", content]],
    }
    section_with_check = _add_checkbox_column(section_as_table)
    table_builder.build(section_with_check, doc)


def _render_note(doc: Document, note: str, style_mgr: StyleManager) -> None:
    """渲染 note 备注文本，灰色小字。"""
    para = doc.add_paragraph()
    run = para.add_run(note)
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
