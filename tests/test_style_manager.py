"""Tests for src/generator/style_manager.py"""

import pytest
from docx import Document
from docx.shared import Pt, RGBColor

from src.generator.style_manager import StyleManager


@pytest.fixture
def sm():
    """Return a StyleManager loaded from config/styles.yaml."""
    return StyleManager()


@pytest.fixture
def doc():
    """Return a blank python-docx Document."""
    return Document()


# ------------------------------------------------------------------ #
# 1. load_styles integration
# ------------------------------------------------------------------ #
class TestLoadStyles:
    def test_styles_loaded(self, sm):
        """StyleManager should have a non-empty styles dict."""
        assert sm.styles
        assert "heading1" in sm.styles
        assert "body" in sm.styles

    def test_heading1_fields(self, sm):
        s = sm.styles["heading1"]
        assert s["font"] == "微软雅黑"
        assert s["size"] == 16
        assert s["bold"] is True
        assert s["color"] == "#1a5276"


# ------------------------------------------------------------------ #
# 2. apply_paragraph_style — headings
# ------------------------------------------------------------------ #
class TestApplyParagraphStyleHeadings:
    @pytest.mark.parametrize(
        "style_name, expected_size, expected_color_hex",
        [
            ("heading1", 16, "1a5276"),
            ("heading2", 14, "2471a3"),
            ("heading3", 12, None),  # heading3 has no color
        ],
    )
    def test_heading_paragraph(self, sm, doc, style_name, expected_size, expected_color_hex):
        para = doc.add_paragraph("Test heading")
        sm.apply_paragraph_style(para, style_name)

        run = para.runs[0]
        assert run.font.name == "微软雅黑"
        assert run.font.size == Pt(expected_size)
        assert run.font.bold is True

        if expected_color_hex:
            r, g, b = (
                int(expected_color_hex[0:2], 16),
                int(expected_color_hex[2:4], 16),
                int(expected_color_hex[4:6], 16),
            )
            assert run.font.color.rgb == RGBColor(r, g, b)


# ------------------------------------------------------------------ #
# 3. apply_paragraph_style — body
# ------------------------------------------------------------------ #
class TestApplyParagraphStyleBody:
    def test_body_paragraph(self, sm, doc):
        para = doc.add_paragraph("Body text")
        sm.apply_paragraph_style(para, style_name="body")

        run = para.runs[0]
        assert run.font.name == "宋体"
        assert run.font.size == Pt(10.5)
        # body has no bold key — should not be set to True
        assert run.font.bold is not True


# ------------------------------------------------------------------ #
# 4. apply_run_style
# ------------------------------------------------------------------ #
class TestApplyRunStyle:
    def test_run_style_heading1(self, sm, doc):
        para = doc.add_paragraph()
        run = para.add_run("Run text")
        sm.apply_run_style(run, "heading1")

        assert run.font.name == "微软雅黑"
        assert run.font.size == Pt(16)
        assert run.font.bold is True
        assert run.font.color.rgb == RGBColor(0x1A, 0x52, 0x76)

    def test_run_style_body(self, sm, doc):
        para = doc.add_paragraph()
        run = para.add_run("Body run")
        sm.apply_run_style(run, "body")

        assert run.font.name == "宋体"
        assert run.font.size == Pt(10.5)


# ------------------------------------------------------------------ #
# 5. apply_cell_style — table_header & table_body
# ------------------------------------------------------------------ #
class TestApplyCellStyle:
    def _make_cell(self, doc, text="cell"):
        table = doc.add_table(rows=1, cols=1)
        cell = table.cell(0, 0)
        cell.text = text
        return cell

    def test_table_header_cell(self, sm, doc):
        cell = self._make_cell(doc, "Header")
        sm.apply_cell_style(cell, "table_header")

        run = cell.paragraphs[0].runs[0]
        assert run.font.name == "微软雅黑"
        assert run.font.size == Pt(10)
        assert run.font.bold is True

        # bg_color shading should be applied (check the XML element exists)
        tc = cell._tc
        shading_elems = tc.findall(
            ".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}shd"
        )
        assert len(shading_elems) >= 1
        fill_val = shading_elems[0].get(
            "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill"
        )
        assert fill_val.upper() == "F2F3F4"

    def test_table_body_cell(self, sm, doc):
        cell = self._make_cell(doc, "Body cell")
        sm.apply_cell_style(cell, "table_body")

        run = cell.paragraphs[0].runs[0]
        assert run.font.name == "宋体"
        assert run.font.size == Pt(10)
        # table_body has no bold — should not be True
        assert run.font.bold is not True

    def test_table_body_no_bg_shading(self, sm, doc):
        """table_body has no bg_color, so no shading element should be added."""
        cell = self._make_cell(doc, "No bg")
        sm.apply_cell_style(cell, "table_body")

        tc = cell._tc
        tc_pr = tc.find(
            "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tcPr"
        )
        if tc_pr is not None:
            shading_elems = tc_pr.findall(
                "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}shd"
            )
            assert len(shading_elems) == 0


# ------------------------------------------------------------------ #
# 6. Missing optional fields don't crash
# ------------------------------------------------------------------ #
class TestMissingOptionalFields:
    def test_apply_style_missing_color(self, sm, doc):
        """heading3 has no 'color' key — should not raise."""
        para = doc.add_paragraph("No color heading")
        sm.apply_paragraph_style(para, "heading3")
        run = para.runs[0]
        assert run.font.name == "微软雅黑"
        assert run.font.size == Pt(12)
        assert run.font.bold is True

    def test_apply_style_missing_bold(self, sm, doc):
        """body has no 'bold' key — should not raise."""
        para = doc.add_paragraph("No bold body")
        sm.apply_paragraph_style(para, "body")

    def test_apply_cell_style_missing_bg_color(self, sm, doc):
        """table_body has no 'bg_color' — should not raise."""
        table = doc.add_table(rows=1, cols=1)
        cell = table.cell(0, 0)
        cell.text = "ok"
        sm.apply_cell_style(cell, "table_body")

    def test_apply_style_with_minimal_config(self, sm, doc):
        """Manually inject a style with only 'font' — should not crash."""
        sm.styles["_minimal"] = {"font": "Arial"}
        para = doc.add_paragraph("Minimal")
        sm.apply_paragraph_style(para, "_minimal")
        run = para.runs[0]
        assert run.font.name == "Arial"
