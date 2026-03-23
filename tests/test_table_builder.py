import pytest
from docx import Document
from src.generator.table_builder import TableBuilder


def test_build_key_value_table():
    doc = Document()
    builder = TableBuilder()
    section = {
        "type": "key_value_table",
        "columns": ["项目", "内容"],
        "rows": [["项目名称", "测试项目"], ["采购编号", "TEST001"]],
    }
    builder.build(section, doc)
    assert len(doc.tables) == 1
    assert doc.tables[0].rows[0].cells[0].text == "项目"


def test_build_standard_table():
    doc = Document()
    builder = TableBuilder()
    section = {
        "type": "standard_table",
        "columns": ["序号", "要求", "说明"],
        "rows": [["1", "营业执照", "提供复印件"]],
    }
    builder.build(section, doc)
    assert len(doc.tables) == 1
    assert len(doc.tables[0].columns) == 3


def test_build_table_empty_rows():
    """Section with columns but empty rows list should still create a table with header row."""
    doc = Document()
    builder = TableBuilder()
    section = {
        "type": "standard_table",
        "columns": ["列A", "列B"],
        "rows": [],
    }
    builder.build(section, doc)
    assert len(doc.tables) == 1
    table = doc.tables[0]
    # Only the header row should exist
    assert len(table.rows) == 1
    assert table.rows[0].cells[0].text == "列A"
    assert table.rows[0].cells[1].text == "列B"


def test_build_multiple_tables():
    """Calling build twice should add two tables to the same document."""
    doc = Document()
    builder = TableBuilder()
    section_a = {
        "type": "key_value_table",
        "columns": ["键", "值"],
        "rows": [["k1", "v1"]],
    }
    section_b = {
        "type": "standard_table",
        "columns": ["A", "B", "C"],
        "rows": [["1", "2", "3"]],
    }
    builder.build(section_a, doc)
    builder.build(section_b, doc)
    assert len(doc.tables) == 2


def test_build_table_section_title():
    """If section has a 'title' field, it should appear as a paragraph before the table."""
    doc = Document()
    builder = TableBuilder()
    section = {
        "type": "key_value_table",
        "title": "基本信息表",
        "columns": ["项目", "内容"],
        "rows": [["名称", "示例"]],
    }
    builder.build(section, doc)
    # The title paragraph should be the first element in the body
    paragraphs = [p for p in doc.paragraphs if p.text.strip()]
    assert len(paragraphs) >= 1
    assert paragraphs[0].text == "基本信息表"
    # Table should still be present
    assert len(doc.tables) == 1
