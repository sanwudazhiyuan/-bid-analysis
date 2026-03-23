"""Tests for checklist generator (资料清单生成器)"""

import os

import pytest
from docx import Document

from src.generator.checklist_gen import render_checklist


@pytest.fixture
def extracted_data():
    """模拟提取结果，包含资料清单模块"""
    return {
        "schema_version": "1.0",
        "modules": {
            "checklist": {
                "title": "资料清单",
                "sections": [
                    {
                        "id": "CL1",
                        "title": "资格证明材料",
                        "type": "standard_table",
                        "columns": ["序号", "材料名称", "要求", "备注"],
                        "rows": [
                            ["1", "营业执照", "复印件加盖公章", "必须"],
                            ["2", "税务登记证", "复印件加盖公章", "必须"],
                            ["3", "社保缴纳证明", "近三个月", "必须"],
                        ],
                    },
                    {
                        "id": "CL2",
                        "title": "技术能力材料",
                        "type": "standard_table",
                        "columns": ["序号", "材料名称", "要求", "备注"],
                        "rows": [
                            ["1", "类似项目业绩", "合同复印件", "至少3个"],
                            ["2", "技术方案", "详细方案书", "必须"],
                        ],
                    },
                ],
            },
        },
    }


def test_render_checklist_basic(tmp_path, extracted_data):
    """生成资料清单 .docx，包含分类表格"""
    out = str(tmp_path / "checklist.docx")
    render_checklist(extracted_data, out)
    assert os.path.exists(out)

    doc = Document(out)
    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert "资料清单" in full_text or "资格证明" in full_text
    assert len(doc.tables) >= 2


def test_render_checklist_table_content(tmp_path, extracted_data):
    """验证表格内容正确"""
    out = str(tmp_path / "checklist.docx")
    render_checklist(extracted_data, out)

    doc = Document(out)
    # 第一个表格应包含资格证明材料
    first_table = doc.tables[0]
    all_text = " ".join(
        cell.text for row in first_table.rows for cell in row.cells
    )
    assert "营业执照" in all_text


def test_render_checklist_no_module(tmp_path):
    """没有 checklist 模块时应生成占位内容"""
    data = {
        "schema_version": "1.0",
        "modules": {
            "module_a": {"title": "A. 项目概况", "sections": []},
        },
    }
    out = str(tmp_path / "checklist.docx")
    render_checklist(data, out)
    assert os.path.exists(out)


def test_render_checklist_failed_module(tmp_path):
    """checklist 失败时应渲染占位文本"""
    data = {
        "schema_version": "1.0",
        "modules": {
            "checklist": {"status": "failed", "error": "提取失败"},
        },
    }
    out = str(tmp_path / "checklist.docx")
    render_checklist(data, out)
    assert os.path.exists(out)


def test_render_checklist_creates_parent_dirs(tmp_path, extracted_data):
    """自动创建输出目录"""
    out = str(tmp_path / "sub" / "checklist.docx")
    render_checklist(extracted_data, out)
    assert os.path.exists(out)
