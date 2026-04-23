"""Tests for bid format generator (投标文件格式生成器)"""

import os

import pytest
from docx import Document

from src.generator.format_gen import render_format


@pytest.fixture
def extracted_data():
    """模拟提取结果，包含投标文件格式相关模块"""
    return {
        "schema_version": "1.0",
        "modules": {
            "bid_format": {
                "title": "投标文件格式",
                "sections": [
                    {
                        "id": "BF1",
                        "title": "投标函",
                        "type": "template",
                        "content": (
                            "致：采购人\n"
                            "根据贵方采购项目（编号：TEST001）的投标邀请，"
                            "我方授权代表签署本投标函。"
                        ),
                    },
                    {
                        "id": "BF2",
                        "title": "报价表",
                        "type": "standard_table",
                        "columns": ["序号", "服务内容", "单价（元）", "数量", "总价（元）"],
                        "rows": [
                            ["1", "制卡服务", "", "", ""],
                            ["2", "递卡服务", "", "", ""],
                        ],
                    },
                    {
                        "id": "BF3",
                        "title": "服务承诺书",
                        "type": "template",
                        "content": "我公司承诺按照招标文件要求提供全部服务。",
                    },
                ],
            },
        },
    }


def test_render_format_basic(tmp_path, extracted_data):
    """生成投标文件格式 .docx，包含投标函和报价表"""
    out = str(tmp_path / "format.docx")
    render_format(extracted_data, out)
    assert os.path.exists(out)

    doc = Document(out)
    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert "投标函" in full_text
    assert "报价表" in full_text or len(doc.tables) >= 1
    assert "服务承诺书" in full_text


def test_render_format_has_table(tmp_path, extracted_data):
    """报价表应生成为表格"""
    out = str(tmp_path / "format.docx")
    render_format(extracted_data, out)

    doc = Document(out)
    assert len(doc.tables) >= 1
    # 报价表有 5 列
    assert len(doc.tables[0].columns) == 5


def test_render_format_no_bid_format_module(tmp_path):
    """没有 bid_format 模块时应生成占位内容，不崩溃"""
    data = {
        "schema_version": "1.0",
        "modules": {
            "module_a": {
                "title": "A. 项目概况",
                "sections": [],
            }
        },
    }
    out = str(tmp_path / "format.docx")
    render_format(data, out)
    assert os.path.exists(out)


def test_render_format_failed_module(tmp_path):
    """bid_format 失败时应渲染占位文本"""
    data = {
        "schema_version": "1.0",
        "modules": {
            "bid_format": {"status": "failed", "error": "提取失败"},
        },
    }
    out = str(tmp_path / "format.docx")
    render_format(data, out)
    assert os.path.exists(out)


def test_render_format_creates_parent_dirs(tmp_path, extracted_data):
    """自动创建输出目录"""
    out = str(tmp_path / "sub" / "format.docx")
    render_format(extracted_data, out)
    assert os.path.exists(out)


def test_render_format_nodes_tree(tmp_path):
    """新结构（nodes 树）渲染：包含编号、动态节点、样表嵌入"""
    data = {
        "schema_version": "1.0",
        "modules": {
            "bid_format": {
                "title": "投标文件",
                "nodes": [
                    {
                        "title": "投标函",
                        "level": 1,
                        "number": "一、",
                        "source": "format_template",
                        "has_sample": True,
                        "dynamic": False,
                        "dynamic_hint": None,
                        "sample_content": {"type": "text", "content": "致：采购人"},
                        "children": [],
                    },
                    {
                        "title": "技术部分",
                        "level": 1,
                        "number": "二、",
                        "source": "scoring_factor",
                        "has_sample": False,
                        "dynamic": False,
                        "dynamic_hint": None,
                        "children": [
                            {
                                "title": "实施方案",
                                "level": 2,
                                "number": "2.1",
                                "source": "scoring_factor",
                                "has_sample": False,
                                "dynamic": True,
                                "dynamic_hint": "根据评分细则展开",
                                "children": [],
                            },
                        ],
                    },
                    {
                        "title": "报价表",
                        "level": 1,
                        "number": "三、",
                        "source": "format_template",
                        "has_sample": True,
                        "dynamic": False,
                        "dynamic_hint": None,
                        "sample_content": {
                            "type": "standard_table",
                            "columns": ["序号", "服务内容", "单价"],
                            "rows": [["1", "制卡服务", "100"]],
                        },
                        "children": [],
                    },
                ],
            },
        },
    }
    out = str(tmp_path / "outline.docx")
    render_format(data, out)
    assert os.path.exists(out)

    doc = Document(out)
    full_text = "\n".join(p.text for p in doc.paragraphs)
    # 目录 + 标题 + 动态提示
    assert "投标函" in full_text
    assert "实施方案" in full_text
    assert "根据评分细则展开" in full_text
    # 样表嵌入 → 表格
    assert len(doc.tables) >= 1
