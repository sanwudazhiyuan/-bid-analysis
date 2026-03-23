"""Tests for report generator (分析报告生成器)"""

import os

import pytest
from docx import Document
from docx.shared import RGBColor

from src.generator.report_gen import render_report


@pytest.fixture
def basic_data():
    return {
        "schema_version": "1.0",
        "modules": {
            "module_a": {
                "title": "A. 项目概况",
                "sections": [
                    {
                        "id": "A1",
                        "title": "基本信息",
                        "type": "key_value_table",
                        "columns": ["项目", "内容"],
                        "rows": [
                            ["项目名称", "测试项目"],
                            ["采购编号", "TEST001"],
                        ],
                    }
                ],
            }
        },
    }


@pytest.fixture
def multi_module_data():
    return {
        "schema_version": "1.0",
        "modules": {
            "module_a": {
                "title": "A. 项目概况",
                "sections": [
                    {
                        "id": "A1",
                        "title": "基本信息",
                        "type": "key_value_table",
                        "columns": ["项目", "内容"],
                        "rows": [["项目名称", "测试项目"]],
                    }
                ],
            },
            "module_b": {
                "title": "B. 资格要求",
                "sections": [
                    {
                        "id": "B1",
                        "title": "资格条件",
                        "type": "standard_table",
                        "columns": ["序号", "要求", "说明"],
                        "rows": [["1", "营业执照", "提供复印件"]],
                    }
                ],
            },
            "module_d": {
                "title": "D. 评分标准",
                "sections": [
                    {
                        "id": "D1",
                        "title": "技术评分",
                        "type": "standard_table",
                        "columns": ["序号", "评分项", "分值"],
                        "rows": [["1", "方案", "30"], ["2", "价格", "40"]],
                    },
                    {
                        "id": "D2",
                        "title": "商务评分",
                        "type": "standard_table",
                        "columns": ["序号", "评分项", "分值"],
                        "rows": [["1", "业绩", "30"]],
                    },
                ],
            },
        },
    }


def test_render_report_basic(tmp_path, basic_data):
    """基本报告生成：输出 .docx 可打开，包含模块标题和表格"""
    out = str(tmp_path / "report.docx")
    render_report(basic_data, out)
    assert os.path.exists(out)

    doc = Document(out)
    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert "A. 项目概况" in full_text
    assert len(doc.tables) >= 1


def test_render_report_tables_have_checkbox_column(tmp_path, basic_data):
    """报告中的表格应包含勾选列"""
    out = str(tmp_path / "report.docx")
    render_report(basic_data, out)

    doc = Document(out)
    table = doc.tables[0]
    # 最后一列应是勾选列
    last_col_header = table.rows[0].cells[-1].text
    assert "确认" in last_col_header or "✓" in last_col_header or "完成" in last_col_header


def test_render_report_multi_module(tmp_path, multi_module_data):
    """多模块报告：所有模块标题和表格均存在"""
    out = str(tmp_path / "report.docx")
    render_report(multi_module_data, out)

    doc = Document(out)
    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert "A. 项目概况" in full_text
    assert "B. 资格要求" in full_text
    assert "D. 评分标准" in full_text
    # D 模块有 2 个 section，加上 A 和 B 各 1 个，共 4 个表格
    assert len(doc.tables) >= 4


def test_render_report_failed_module_placeholder(tmp_path):
    """失败模块应渲染红色占位文本，不应导致生成器崩溃"""
    data = {
        "schema_version": "1.0",
        "modules": {
            "module_a": {
                "title": "A. 项目概况",
                "sections": [
                    {
                        "id": "A1",
                        "title": "基本信息",
                        "type": "key_value_table",
                        "columns": ["项目", "内容"],
                        "rows": [["项目名称", "测试项目"]],
                    }
                ],
            },
            "module_c": {"status": "failed", "error": "LLM 返回非法 JSON"},
        },
    }
    out = str(tmp_path / "report_with_failure.docx")
    render_report(data, out)
    assert os.path.exists(out)

    doc = Document(out)
    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert "提取失败" in full_text or "failed" in full_text.lower()

    # 验证占位文本为红色
    found_red = False
    for para in doc.paragraphs:
        for run in para.runs:
            if "提取失败" in run.text or "failed" in run.text.lower():
                if run.font.color.rgb == RGBColor(0xFF, 0x00, 0x00):
                    found_red = True
    assert found_red, "失败模块的占位文本应为红色"


def test_render_report_none_module(tmp_path):
    """None 模块应渲染占位文本，不崩溃"""
    data = {
        "schema_version": "1.0",
        "modules": {
            "module_a": None,
            "module_b": {
                "title": "B. 资格要求",
                "sections": [],
            },
        },
    }
    out = str(tmp_path / "report_none.docx")
    render_report(data, out)
    assert os.path.exists(out)

    doc = Document(out)
    full_text = "\n".join(p.text for p in doc.paragraphs)
    # 应包含某种占位信息
    assert "无数据" in full_text or "module_a" in full_text.lower() or "跳过" in full_text


def test_render_report_empty_sections(tmp_path):
    """模块无 sections 时不崩溃"""
    data = {
        "schema_version": "1.0",
        "modules": {
            "module_a": {
                "title": "A. 项目概况",
                "sections": [],
            },
        },
    }
    out = str(tmp_path / "report_empty.docx")
    render_report(data, out)
    assert os.path.exists(out)

    doc = Document(out)
    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert "A. 项目概况" in full_text


def test_render_report_creates_parent_dirs(tmp_path):
    """输出路径的父目录不存在时应自动创建"""
    out = str(tmp_path / "sub" / "dir" / "report.docx")
    data = {
        "schema_version": "1.0",
        "modules": {
            "module_a": {
                "title": "A. 项目概况",
                "sections": [],
            },
        },
    }
    render_report(data, out)
    assert os.path.exists(out)


def test_render_report_excludes_bid_format_and_checklist(tmp_path):
    """分析报告不应包含 bid_format 和 checklist 模块"""
    data = {
        "schema_version": "1.0",
        "modules": {
            "module_a": {
                "title": "A. 项目概况",
                "sections": [
                    {
                        "id": "A1",
                        "title": "基本信息",
                        "type": "key_value_table",
                        "columns": ["项目", "内容"],
                        "rows": [["项目名称", "测试"]],
                    }
                ],
            },
            "bid_format": {
                "title": "投标文件格式",
                "sections": [
                    {
                        "id": "BF1",
                        "title": "投标函",
                        "type": "template",
                        "content": "这段内容不应出现在分析报告中",
                    }
                ],
            },
            "checklist": {
                "title": "资料清单",
                "sections": [
                    {
                        "id": "CL1",
                        "title": "资格材料",
                        "type": "standard_table",
                        "columns": ["序号", "材料"],
                        "rows": [["1", "营业执照"]],
                    }
                ],
            },
        },
    }
    out = str(tmp_path / "report.docx")
    render_report(data, out)

    doc = Document(out)
    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert "A. 项目概况" in full_text
    assert "投标文件格式" not in full_text
    assert "投标函" not in full_text
    assert "资料清单" not in full_text


def test_render_report_section_numbering(tmp_path):
    """子标题应使用编号格式如 A.1, A.2"""
    data = {
        "schema_version": "1.0",
        "modules": {
            "module_a": {
                "title": "A. 项目概况",
                "sections": [
                    {
                        "id": "A1",
                        "title": "基本信息",
                        "type": "key_value_table",
                        "columns": ["项目", "内容"],
                        "rows": [["项目名称", "测试"]],
                    },
                    {
                        "id": "A2",
                        "title": "采购信息",
                        "type": "key_value_table",
                        "columns": ["项目", "内容"],
                        "rows": [["预算", "100万"]],
                    },
                ],
            },
        },
    }
    out = str(tmp_path / "report.docx")
    render_report(data, out)

    doc = Document(out)
    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert "A.1" in full_text
    assert "A.2" in full_text


def test_render_report_text_sections_converted_to_table(tmp_path):
    """text 类型的 section 应转为单列表格而非文字段落"""
    data = {
        "schema_version": "1.0",
        "modules": {
            "module_d": {
                "title": "D. 合同条款",
                "sections": [
                    {
                        "id": "D1",
                        "title": "知识产权",
                        "type": "text",
                        "content": "甲方拥有银行卡的设计权和知识产权。",
                    },
                    {
                        "id": "D2",
                        "title": "保密条款",
                        "type": "text",
                        "content": "甲乙双方均负有保密义务。",
                    },
                ],
            },
        },
    }
    out = str(tmp_path / "report.docx")
    render_report(data, out)

    doc = Document(out)
    # text 内容应在表格中，不应作为独立段落
    all_para_text = "\n".join(p.text for p in doc.paragraphs)
    assert "甲方拥有银行卡" not in all_para_text
    assert "甲乙双方" not in all_para_text
    # 应有表格包含这些内容
    assert len(doc.tables) >= 2
    all_table_text = " ".join(
        cell.text for t in doc.tables for row in t.rows for cell in row.cells
    )
    assert "甲方拥有银行卡" in all_table_text
    assert "甲乙双方" in all_table_text
